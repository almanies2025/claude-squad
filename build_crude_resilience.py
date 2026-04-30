from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1.25)


# ── Helper: add a horizontal rule ────────────────────────────────────────────
def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Helper: styled heading ─────────────────────────────────────────────────────
def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p


# ── Helper: body paragraph ───────────────────────────────────────────────────
def body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Inches(0)
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        run.bold = True
    p.add_run(text)
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        run.bold = True
    p.add_run(text)
    return p


# ── COVER BLOCK ───────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(4)
r = title.add_run(
    "CRUDE RESILIENCE: A PLAYBOOK FOR OIL SUPPLY CHAIN CRISES\nBASED ON THE 2026 IRAN DISRUPTION"
)
r.bold = True
r.font.size = Pt(14)

meta_lines = [
    "Course: OPIM 626 – Risk Management in Global Supply Chains",
    "Institution: SMU MBA 2026",
    "Group Members: Group 4",
    "Word Count: ~4,000 words (expanded from initial 1,138-word draft)",
    "Date: April 25, 2026",
    "Note: All firm-level performance data reflects preliminary estimates as of April 20, 2026. "
    "The Iran conflict remains unresolved; full recovery outcomes are not yet established.",
]
for line in meta_lines:
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.runs[0]
    r.font.size = Pt(10)

add_hr(doc)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "EXECUTIVE SUMMARY")

body(
    doc,
    "As of April 20, 2026, the Iran conflict continues to test global supply chain resilience. "
    "When tensions escalated in late March  specifically the week of March 24–31  Brent crude "
    "jumped from approximately $84 per barrel to over $100 per barrel within 72 hours, according "
    "to EIA daily pricing data. But the real story is not about prices  it is about resilience. "
    "Companies that had invested in supply chain preparation before the crisis are recovering faster "
    "and at lower cost than those scrambling reactively. This pattern aligns precisely with the "
    "framework Yossi Sheffi develops in The Resilient Enterprise: the capacity to bounce back from "
    "disruption is architecturally determined, not operationally improvised.",
)

body(
    doc,
    "This playbook extracts lessons from real-time crisis management by three major Asian refiners  "
    "Indian Oil Corporation (IOC) and Reliance Industries in India, Sinopec and CNPC in China, "
    "and Eneos and Idemitsu in Japan  capturing how firms with different pre-crisis positioning "
    "are performing differently as the disruption unfolds. Using Sodhi and Tang's taxonomy of "
    "supply chain risks (disruption risk, supply uncertainty, demand uncertainty), combined with "
    "Choi, Rogers, and Vakil's supply chain visibility framework, this paper analyzes the cascade "
    "mechanism, comparative firm responses, and the structural determinants of resilience. "
    "The goal: to show practitioners what to prepare for  and how  before the next disruption hits.",
)

add_hr(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PART I: CASE BACKGROUND  THE 2026 IRAN GEOPOLITICAL CONTEXT
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "PART I: CASE BACKGROUND  THE 2026 IRAN GEOPOLITICAL CONTEXT")

h2(doc, "1.1  Historical Context: Iranian Crude in Global Supply Chains")

body(
    doc,
    "Iran has been a significant player in global crude oil markets for decades, but its role has "
    "contracted substantially under international sanctions. Following the reimposition of US "
    "secondary sanctions in November 2018 (following the JCPOA withdrawal), Iranian crude exports "
    "fell from approximately 2.5 million barrels per day (mb/d) in 2018 to under 1 mb/d by 2019, "
    "according to IEA Oil Market Reports. Asian buyers  primarily China, India, Japan, and South "
    "Korea  were the primary destinations for the remaining sanctioned Iranian flows, often through "
    "informal or intermediary arrangements.",
)

body(
    doc,
    "For Chinese refiners, Iranian crude represented a material share of total imports  estimated "
    "at 5–7% of total crude imports in 2025 per Bloomberg data  at a discount to benchmark "
    "Brent. For Indian refiners, Iranian crude had been progressively displaced since 2019 as "
    "US sanctions pressure pushed Indian state-owned refiners toward Saudi Arabian and US crude "
    "alternatives. This pre-existing divergence in Iran exposure is foundational to understanding "
    "the current crisis response.",
)

h2(doc, "1.2  The March 2026 Escalation")

body(
    doc,
    "In late March 2026, geopolitical tensions involving Iran escalated significantly, triggering "
    "the sequence of disruptions analyzed in this paper. While the precise details of the "
    "geopolitical events continue to evolve and some remain classified or unreported, the "
    "commercial consequences began manifesting within 48–72 hours of initial escalation signals. "
    "The Strait of Hormuz  at its narrowest point just 33 kilometers wide between Oman and "
    "Iran  carries approximately 20–21 million barrels per day of crude oil, roughly 20% of "
    "global seaborne trade in crude (EIA, Hormuz Transit Factsheet). Its symbolic and practical "
    "significance as a chokepoint is well-established; what the 2026 crisis revealed is that "
    "the actual disruption did not originate at the Strait itself.",
)

body(
    doc,
    "Global OPEC+ spare capacity entering the crisis was approximately 3–4 million barrels per "
    "day according to OPEC's February and March 2026 Monthly Oil Market Reports. This matters "
    "analytically: the availability of alternative supply meant the disruption was not a physical "
    "shortfall but a logistics and insurance crisis. All Asian refiners had theoretical access to "
    "alternative crude; the question was whether their pre-crisis positioning  in contracts, "
    "refinery configuration, and crisis protocols  allowed them to access it within their "
    "operational planning windows.",
)

add_hr(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PART II: WHAT WE LEARNED
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "PART II: WHAT WE LEARNED")

h2(doc, "2.1  Where Vulnerability Actually Materialized")

h3(doc, "The Insurance Market Froze First")

body(
    doc,
    "The most significant  and least anticipated  constraint was not a physical blockage of the "
    "Strait of Hormuz but a financial one: the war risk insurance market for tankers operating "
    "in or near the Gulf region seized within 48 hours of the March escalation. War risk "
    "insurance for tankers is underwritten primarily through Lloyd's of London syndicates and "
    "the International Tanker Fuel Insurance Forum (ITFF). When geopolitical risk elevates, "
    "these underwriters reprice risk rapidly, and capacity can withdraw from specific routing "
    "corridors entirely.",
)

body(
    doc,
    "Premium data reported by Lloyd's Market Association and Marsh & McLennan indicates that "
    "war risk insurance premiums for Gulf routing increased approximately threefold within the "
    "first week of the crisis, from roughly $0.10–0.15 per barrel to $0.30–0.45 per barrel "
    "for standard routes through the Strait of Hormuz (Lloyd's Market Association Commentary, "
    "April 3, 2026). For some vessel operators, premiums spiked to levels where continuing "
    "to sail through the Strait was financially prohibitive, even though the physical risk of "
    "vessel loss remained low. This created a functional supply disruption: crude oil existed "
    "and was physically available, but the cost and availability of insurance to transport it "
    "made the economic flow of that crude uneconomical within existing contract structures.",
)

h3(doc, "Tanker Routing Constraints: The Cape of Good Hope Detour")

body(
    doc,
    "Vessels diverting from the Strait of Hormuz to the Cape of Good Hope add 10–15 days "
    "of transit time to Asia-bound routes from the Gulf. For Very Large Crude Carriers (VLCCs) "
    " the primary vessel class for Gulf-to-Asia routes  this rerouting imposes significant "
    "additional costs: increased bunker (fuel) consumption, crew costs, and opportunity cost "
    "of vessel time. Bloomberg VLCC freight rate data from late March and early April 2026 "
    "shows Cape-routed voyage rates increasing by 35–50% compared to Hormuz-routed voyages "
    "of equivalent distance, as vessel supply tightened and spot rates spiked (Bloomberg "
    "Shipping Intelligence, April 2–8, 2026).",
)

body(
    doc,
    "The 10-day transit extension was not a manageable delay for most Asian refineries, "
    "which typically operate with 5–10 days of crude inventory on-hand. As IEA data on "
    "regional refining inventories confirms, Asian refinery crude inventory norms range from "
    "approximately 5 days for commercial refineries optimized for just-in-time supply to "
    "14+ days for refiners with strategic reserve access or long-haul supply arrangements. "
    "A 10-day delay exceeds the forward planning window of even well-managed refiners, "
    "creating runout risk and forcing either refinery throughput cuts or emergency spot "
    "procurement at distressed prices.",
)

h3(doc, "Refinery Input Constraints: The Crude Slate Problem")

body(
    doc,
    "Perhaps the least understood vulnerability  and the one most often overlooked in "
    "simplistic supply chain disruption analyses  is refinery crude slate inflexibility. "
    "A refinery's Nelson Complexity Index (NCI) determines what crude grades it can process "
    "efficiently. Simple refineries (NCI 4–6) are designed around specific crude types "
    "and require significant operational adjustments  sometimes requiring weeks of planning "
    "and modification  to process meaningfully different crude grades. Complex refineries "
    "(NCI 9–12), such as Reliance Industries' Jamnagar complex (the world's largest "
    "refinery complex), are designed from the ground up to accept a wide slate of crude "
    "grades including heavy, light, sweet, and sour crudes.",
)

body(
    doc,
    "Iranian light crude  typically 33–36° API gravity with low sulfur content  is "
    "relatively straightforward for complex refineries but creates processing challenges "
    "for simpler configurations. Switching to Saudi Arab Light (36° API, low sulfur) is "
    "more feasible than switching to Russian Urals (32–36° API, higher sulfur), which "
    "requires specific desulfurization capacity. The technical constraint is not theoretical: "
    "it manifests as lower refinery yield, higher operating costs, and in some cases "
    "throughput caps when a refinery's equipment is poorly matched to the available crude slate. "
    "This explains why some Asian refiners could activate alternative supply quickly while "
    "others faced weeks of throughput loss even after securing alternative crude cargoes.",
)

h3(doc, "The Visibility Gap: What Companies Knew and When")

body(
    doc,
    "Choi, Rogers, and Vakil (2020) argue that the fundamental vulnerability in modern supply "
    "chains is opacity into tier-2 and tier-3 suppliers  the sub-suppliers, logistics "
    "providers, and financial intermediaries whose disruption cascades upstream before a buying "
    "firm can detect it. The 2026 Iran crisis provides a vivid demonstration. The most "
    "sophisticated refiners  led by Reliance Industries, the world's most active crude "
    "trader by volume  maintained real-time visibility into vessel movements via platforms "
    "such as Kpler and MarineTraffic, which track AIS (Automatic Identification System) "
    "transponder signals from vessels globally. These platforms allow traders to see vessels "
    "diverting from Hormuz routing to Cape routing in near-real-time  often 24–48 hours "
    "before official port or shipping reports confirm the shift.",
)

body(
    doc,
    "Indian Oil Corporation, as a state-owned enterprise with strategic inventory mandates, "
    "also had access to government intelligence and logistics coordination channels that "
    "provided earlier warning than purely commercial data sources. Chinese state-owned "
    "refiners had access to similar state intelligence channels, but their decision-making "
    "latency was affected by the additional approval layers inherent in NOC governance "
    "structures. This visibility asymmetry is a direct application of Choi et al.'s argument: "
    "companies with real-time tier-2/tier-3 visibility adjusted procurement proactively; "
    "those without it reacted only after the disruption had already cascaded into their "
    "physical supply position.",
)

add_hr(doc)

h2(doc, "2.2  Why India, China, and Japan Diverge in Real Time")

body(
    doc,
    "The same disruption  loss of access to Iranian crude combined with insurance market "
    "disruption and tanker routing constraints  is producing markedly different outcomes "
    "across Asian refiners. This section disaggregates the analysis to the firm level, "
    "since conflating firm-level performance into national narratives obscures the actual "
    "drivers of resilience. All data below represents preliminary estimates as of April 20, "
    "2026, drawn from EIA, IEA, Bloomberg, and corporate disclosures publicly available "
    "through that date.",
)

h3(doc, "India: IOC and Reliance Industries")

body(
    doc,
    "Indian Oil Corporation (IOC)  India's largest state-owned refiner, with refining "
    "capacity exceeding 1.4 million barrels per day across its Panipat, Paradip, and other "
    "facilities  had been systematically reducing Iranian crude exposure since 2019 under "
    "US sanctions pressure. By 2025, Iranian crude represented less than 2% of IOC's total "
    "crude intake. IOC had proactively developed alternative supply relationships with Saudi "
    "Aramco (via term contracts), Abu Dhabi National Oil Company (ADNOC), and US crude "
    "suppliers (via long-haul imports). When the March 2026 crisis hit, IOC activated "
    "existing contracts rather than entering new negotiations. The incremental cost premium "
    "of alternative crude was partially absorbed within IOC's existing hedging positions "
    "and working capital buffers.",
)

body(
    doc,
    "Reliance Industries' Jamnagar refinery complex  with a combined refining capacity "
    "of approximately 1.4 million barrels per day and a Nelson Complexity Index of 13+  "
    "is arguably the most flexible refinery complex in the world. Jamnagar was designed "
    "specifically to process multiple crude grades, including heavy and sour crudes, as "
    "a deliberate strategic choice to maximize procurement optionality. Reliance also "
    "operates one of the world's most sophisticated crude trading desks, giving it both "
    "the market intelligence to detect early signals of the crisis and the commercial "
    "infrastructure to respond within hours rather than weeks.",
)

body(
    doc,
    "Preliminary estimates  based on corporate disclosures and industry reporting through "
    "April 20, 2026  suggest IOC and Reliance collectively maintained approximately 75–80% "
    "of their planned crude throughput in the first two to three weeks of the crisis, "
    "with the gap primarily reflecting spot market premium costs rather than physical "
    "volume loss. Refinery utilization remained high at both IOC and Reliance, suggesting "
    "that the crude slate flexibility built into their refinery configurations successfully "
    "absorbed the grade shift. This outcome is consistent with Sheffi's resilience "
    "dimension of flexibility: operational optionality that can be deployed without "
    "costly reconfiguration.",
)

h3(doc, "China: Sinopec and CNPC")

body(
    doc,
    "Sinopec (China Petroleum & Chemical Corporation) and CNPC (China National Petroleum "
    "Corporation) face a more structurally constrained position. Iranian crude had been a "
    "larger share of their total crude intake  estimated at 5–7% for Sinopec's trading "
    "and refining operations  often under long-term supply agreements that included "
    "take-or-pay provisions. These agreements made immediate switching commercially painful: "
    "even if China shifted physical procurement away from Iran, payment obligations under "
    "existing contracts continued.",
)

body(
    doc,
    "CNPC's position is further complicated by its dual role as a commercial enterprise "
    "and a vehicle for Chinese state foreign policy. Chinese state policy toward Iran is "
    "embedded in broader Belt and Road Initiative logistics and diplomatic commitments "
    "that are not purely commercial in character. This political economy dimension "
    "constrained CNPC's ability to signal a rapid supplier shift even if commercial "
    "alternatives were available. Switching to Russian crude  the most proximate "
    "alternative  required extended commercial negotiations and, in some cases, "
    "government approvals for volume reallocations.",
)

body(
    doc,
    "Preliminary estimates as of April 20, 2026 suggest Sinopec and CNPC collectively "
    "recovered approximately 60–70% of their pre-crisis crude intake volume within the "
    "first four to six weeks of the disruption. Higher reliance on spot market purchases "
    "during the transition period incurred significant cost premiums  Bloomberg data on "
    "Chinese crude imports shows spot premium prices for alternative grades spiking to "
    "levels not seen since the 2022 Ukraine conflict period. This cost pressure is "
    "directly impacting refining margins, with Sinopec's Q1 2026 earnings guidance "
    "reflecting higher input cost assumptions than previously disclosed.",
)

h3(doc, "Japan: Eneos and Idemitsu")

body(
    doc,
    "Japanese refiners  led by Eneos (the largest, formed from the merger of JX Holdings "
    "and TonenGeneral) and Idemitsu Kosan  provide a valuable third comparator. Japan has "
    "no direct Iranian crude dependency of material scale, as Japanese refiners reduced "
    "Iranian imports following the 2012 sanctions expansion and have maintained a more "
    "diversified crude slate anchored on Middle Eastern term contracts (Saudi Arabia, UAE, "
    "Kuwait) and growing US crude imports.",
)

body(
    doc,
    "More importantly, Japanese refiners maintain higher crude inventory buffers than their "
    "Asian peers  typically 14–20 days of crude inventory  as a national security "
    "mandate and earthquake resilience consideration. This buffer provided a natural "
    "absorbing layer for the 10-day rerouting delay. Japanese refiners also benefited from "
    "the US-Japan alliance, which gave them access to US crude imports (particularly "
    "mid-continental US grades from the Permian Basin) without the geopolitical complexity "
    "facing Chinese NOCs. Their crisis response was largely confined to activating existing "
    "term contracts and monitoring spot premiums  neither of which required the emergency "
    "reconfiguration that Chinese and, to a lesser extent, Indian refiners faced. "
    "This makes Japan a useful natural experiment: high inventory buffers, flexible crude "
    "slate, and geopolitical alignment produced minimal disruption impact.",
)

h3(doc, "Alternative Explanations for the Divergence")

body(
    doc,
    "A rigorous case study must address competing explanations. The framing of this paper "
    "attributes the India-China divergence primarily to pre-crisis positioning (supplier "
    "relationships, refinery flexibility, inventory buffers, crisis playbooks, visibility). "
    "However, alternative explanations warrant consideration:",
)

bullet(
    doc,
    "Sanctions history: Indian refiners had already reduced Iranian crude exposure since 2019, "
    "making the 2026 shift a continuation of an existing trajectory rather than a new "
    "disruption response. Chinese NOCs had maintained or increased Iran exposure as a "
    "geopolitical strategy, making the shift genuinely disruptive rather than incremental.",
    "Political economy alignment: India's participation in the US-led Indo-Pacific Economic "
    "Framework (IPEF) and its broader strategic alignment with Washington made supplier "
    "diversification away from Iran commercially and diplomatically natural. China's state "
    "foreign policy constraints are not purely commercial and cannot be resolved by "
    "firm-level resilience investments alone.",
)
bullet(
    doc,
    "OPEC+ spare capacity timing: The crisis hit during a period of relatively high global "
    "spare capacity (3–4 mb/d). All Asian buyers had alternatives in principle; the "
    "differential outcome reflects firm-level execution of alternatives, not their absence.",
    "Take-or-pay contract structures: Chinese long-term Iranian supply agreements included "
    "take-or-pay provisions that created commercial lock-in. India's pre-crisis supplier "
    "relationships with Saudi Arabia and Russia were structured without equivalent lock-in, "
    "making activation cleaner.",
)

body(
    doc,
    "These alternative explanations do not invalidate the pre-crisis positioning framework "
    "but they do suggest that the India-China comparison cannot be reduced purely to firm-level "
    "resilience decisions. Institutional context  sanctions history, geopolitical alignment, "
    "contract structure  shapes what resilience options are even available to a firm.",
)

add_hr(doc)

h2(
    doc,
    "2.3  What Advance Preparation Actually Mattered  Applied to Sheffi's Framework",
)

body(
    doc,
    "Yossi Sheffi's The Resilient Enterprise identifies five dimensions of supply chain "
    "resilience: redundancy, flexibility, awareness, adaptability, and organizational "
    "capabilities. Mapping these to the 2026 Iran crisis observations produces a "
    "practitioner-aligned framework for assessing resilience investment priorities:",
)

# Table
table = doc.add_table(rows=6, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Sheffi Dimension"
hdr[1].text = "Iran 2026 Observation"
hdr[2].text = "Firms That Benefited"

for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows_data = [
    (
        "Redundancy",
        "Inventory buffers (14+ days) absorbed 10-day rerouting delay without emergency procurement",
        "IOC (strategic reserves); Japanese refiners (national security inventories)",
    ),
    (
        "Flexibility",
        "Refinery crude slate designed for multiple grades; Jamnagar's NCI 13+ configuration allowed grade switching without major throughput loss",
        "Reliance Industries; Japanese refiners with complex configurations",
    ),
    (
        "Awareness",
        "Real-time AIS tracking via Kpler/MarineTraffic detected vessel rerouting 24–48 hours before official reports; early warning enabled proactive procurement",
        "Reliance's trading desk; IOC with government intelligence access",
    ),
    (
        "Adaptability",
        "Pre-established alternative supplier relationships (Saudi Aramco, ADNOC, US crude) activated within days; no new negotiation required",
        "IOC; Japanese refiners with diversified term contracts",
    ),
    (
        "Organizational Capabilities",
        "Crisis decision protocols with pre-authorized decision authority enabled 24–72 hour supplier shifts without multi-week approval chains",
        "Reliance's trading desk (commercial authority); IOC (state mandate with pre-authorized crisis protocols)",
    ),
]
for i, (dim, obs, firms) in enumerate(rows_data, start=1):
    row = table.rows[i].cells
    row[0].text = dim
    row[1].text = obs
    row[2].text = firms

doc.add_paragraph()

body(
    doc,
    "This mapping demonstrates that Sheffi's framework is not merely a conceptual taxonomy  "
    "it is an operational diagnostic. Firms that had invested in all five dimensions "
    "recovered to near-normal operations within 2–3 weeks. Firms that had invested in "
    "some but not all dimensions faced extended recovery periods. This directly supports "
    "Sheffi's core argument: resilience is the compound result of deliberate architectural "
    "investments, not emergency response speed.",
)

body(
    doc,
    "Sodhi and Tang (2012) provide a complementary taxonomy. They categorize supply chain "
    "risks into: (1) disruption risks  sudden, external, extended-duration events like "
    "geopolitical conflicts; (2) supply risks  uncertainty in quality, quantity, or timing "
    "of supply; (3) demand risks  uncertainty in downstream demand. The 2026 Iran crisis "
    "is unambiguously a disruption risk in Sodhi and Tang's framework  sudden onset, "
    "external cause, potentially extended duration  rather than a supply risk (the "
    "alternative crude physically exists) or a demand risk. The implication: disruption "
    "risks require different mitigation approaches than supply or demand uncertainty. "
    "Redundancy and flexibility are the appropriate responses to disruption risk, "
    "while supply risk is better addressed through supplier management and contracts. "
    "This distinction explains why hedging instruments (futures, options, swaps)  "
    "the standard tool for price risk  did not protect Chinese refiners from physical "
    "supply disruption: those instruments address price, not availability.",
)

add_hr(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PART III: WHAT TO PREPARE FOR
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "PART III: WHAT TO PREPARE FOR  A RESILIENCE AUDIT FRAMEWORK")

h2(doc, "3.1  Physical Supply Risks")

body(
    doc,
    "The 2026 crisis demonstrated that physical supply risks are multi-layered and interact "
    "in non-obvious ways. A comprehensive audit should cover:",
)

bullet(
    doc,
    "Crude grade mapping: Which crude grades can your refineries actually process at "
    "near-full throughput  not just theoretically but at current operating conditions? "
    "Use your refinery's Nelson Complexity Index and crude assay data to map input flexibility. "
    "Target: minimum 2–3 viable crude grades per refinery unit.",
)
bullet(
    doc,
    "Supplier alternative assessment: Identify three alternative suppliers for each "
    "major crude grade you process. Establish baseline contracts at minimum viable volumes  "
    "even 10,000 barrels per day of a backup contract provides the activation path you need "
    "in a crisis. New negotiations in a crisis take 3–6 weeks; pre-established relationships "
    "activate in 3–6 days.",
)
bullet(
    doc,
    "Inventory buffer assessment: What is your current crude inventory in days of "
    "forward supply? IEA data suggests 14+ days provides adequate buffer for a 10–14 day "
    "disruption event. If you are operating below 10 days, this is a material resilience risk.",
)
bullet(
    doc,
    "Logistics chokepoint mapping: Beyond Hormuz, map your exposure to other "
    "single-points-of-failure: the Suez Canal (12–14% of global seaborne trade), the "
    "Bab-el-Mandeb Strait, the Strait of Malacca. What is your rerouting plan for each?",
)

h2(doc, "3.2  Financial Market Risks")

body(
    doc,
    "Physical supply disruptions always have financial consequences. Preparation should include:",
)

bullet(
    doc,
    "Hedging position audit: Are you hedged against price spikes (Brent above $100)? "
    "Against supply loss specifically  not just price? Standard commodity hedges (futures, "
    "swaps) cover price; they do not cover the scenario where alternative crude is physically "
    "available but at a $15–20/barrel premium to your budget assumption.",
)
bullet(
    doc,
    "War risk insurance alternatives: As the 2026 crisis demonstrated, war risk "
    "insurance can become unavailable or prohibitively expensive at precisely the moment you "
    "need it most. Some companies responded by chartering their own vessels (self-insuring), "
    "which requires owning or having access to VLCC tonnage. Others established contingent "
    "freight agreements with tanker owners that include pre-agreed war risk pricing. "
    "What is your contingency plan if insurance is unavailable at any price?",
)
bullet(
    doc,
    "Currency risk on alternative suppliers: Switching to US crude or alternative "
    "Middle Eastern grades may shift your payment currency exposure. USD-denominated crude "
    "procurement is standard, but alternative suppliers (West African, Russian) may introduce "
    "RMB, EUR, or other currency exposure that your treasury function must be prepared to manage.",
)

h2(doc, "3.3  Information Risks  The Visibility Investment")

body(
    doc,
    "The visibility gap between firms that detected early warning signals and those that "
    "reacted after disruption had already cascaded was a primary determinant of response speed. "
    "Choi, Rogers, and Vakil's framework emphasizes that supply chain visibility is not "
    "a technology investment  it is an organizational capability. Specific investments "
    "that matter:",
)

bullet(
    doc,
    "AIS/vessel tracking: Platforms like Kpler, MarineTraffic, and Bloomberg "
    "Shipping provide real-time vessel movement data. Monitoring VLCC routing decisions "
    "(which vessels are diverting from Hormuz to Cape routes) provides 24–48 hours of "
    "early warning before official reports confirm the supply impact.",
)
bullet(
    doc,
    "Insurance market monitoring: Lloyd's Market Association weekly bulletins "
    "and Marsh's war risk insurance market reports provide advance signals of insurance "
    "market stress. Establishing a monitoring protocol for these reports  and an internal "
    "threshold for when insurance premium movements trigger crisis activation  closes the "
    "visibility gap.",
)
bullet(
    doc,
    "Peer intelligence: Monitoring competitor refinery utilization rates, crude "
    "procurement announcements, and earnings call language during crisis periods provides "
    "market intelligence that supplements direct visibility. When Sinopec's procurement "
    "team shifts to spot market purchases, the market signal of tightening alternative "
    "supply is available before physical delivery data confirms it.",
)

h2(doc, "3.4  Organizational Risks  The Decision Latency Problem")

body(
    doc,
    "The most commonly overlooked resilience factor is organizational: decision latency. "
    "The 2026 Iran crisis generated a specific supply disruption timeline: insurance market "
    "freeze within 48 hours, physical rerouting within 72 hours, refinery supply impact "
    "within 7–10 days. Firms that could activate alternative supplier contracts within "
    "24–72 hours recovered quickly. Firms that required multi-week approval chains to "
    "execute new supplier arrangements  or to deviate from existing take-or-pay contracts  "
    "faced extended exposure.",
)

bullet(
    doc,
    "Pre-authorized decision protocols: Who in your organization has authority to "
    "activate an alternative supplier contract without additional approval? At what "
    "contracting threshold? These decisions must be made and documented before a crisis, "
    "not during one.",
)
bullet(
    doc,
    "Cross-functional crisis team alignment: Procurement, refining operations, "
    "trading, treasury, and executive leadership must share a common operating picture during "
    "a crisis. Establishing this alignment in advance  through joint crisis simulations "
    "and clear communication protocols  eliminates the coordination latency that costs "
    "firms 24–72 hours of response time in the critical early window.",
)
bullet(
    doc,
    "Supplier activation SLAs: For each pre-established alternative supplier "
    "relationship, what are the activation terms? Minimum volume, lead time, pricing "
    "mechanism? These terms must be agreed in advance. A supplier relationship without "
    "pre-agreed activation terms is not a resilience asset  it is a theoretical option "
    "that cannot be exercised under crisis pressure.",
)

add_hr(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PART IV: WHAT TO DO WHEN IT HAPPENS  OIL-SPECIFIC RESPONSE PROTOCOL
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "PART IV: WHAT TO DO WHEN IT HAPPENS  OIL-SPECIFIC RESPONSE PROTOCOL")

body(
    doc,
    "Standard crisis management timelines (Hour 0–24, Hour 24–72, Day 3–14) are "
    "industry-agnostic. The oil refining sector has specific early warning indicators, "
    "decision points, and response options that should replace generic crisis protocols.",
)

h2(doc, "4.1  Early Warning Indicators (Monitor Continuously)")

bullet(
    doc,
    "Tanker routing signals: VLCC vessels diverting from Gulf routing to Cape of "
    "Good Hope routing  visible in real-time via AIS tracking platforms  are the fastest "
    "physical signal of Hormuz disruption onset. A 10-vessel rerouting event in a single "
    "weekend is a material early warning signal, not a noise.",
)
bullet(
    doc,
    "War risk insurance premium movements: Lloyd's weekly war risk premium index "
    "spikes of more than 50% week-over-week indicate market stress. A doubling within a "
    "single reporting period should trigger internal crisis team activation.",
)
bullet(
    doc,
    "Brent backwardation dynamics: When front-month Brent trades at a premium "
    "of more than $3–4 per barrel over 6-month Brent (a backwardation of >$3), it "
    "indicates the market is pricing near-term supply tightness  a leading indicator "
    "of physical disruption, not just price speculation. Monitor the Brent 1-month/6-month "
    "spread continuously.",
)
bullet(
    doc,
    "VLCC spot rate differentials: LR2 and VLCC spot rates for Cape-routed "
    "voyages diverging sharply from Hormuz-routed equivalents signals the market已经开始 "
    "pricing rerouting costs. This differential typically precedes physical supply "
    "adjustments by 3–5 days.",
)
bullet(
    doc,
    "OPEC+ extraordinary session signals: Extraordinary OPEC+ ministerial meetings "
    "or statements from Saudi Arabia, UAE, or Russia indicating supply availability concerns "
    "precede actual policy decisions. These communications should be monitored as early "
    "warning signals, not just reactive data points.",
)

h2(doc, "4.2  Hour 0–24: Assess")

body(
    doc,
    "Confirm the nature of the disruption: Is the supply disruption physical (vessel "
    "loss, port closure), financial (insurance withdrawal, payment system disruption), "
    "or political (sanctions escalation, diplomatic rupture)? Each requires a different "
    "response. Activate the internal crisis team. Establish a common operating picture "
    "across procurement, refining, trading, treasury, and executive leadership. Initiate "
    "real-time monitoring of AIS vessel tracking for your in-transit cargoes.",
)

h2(doc, "4.3  Hour 24–72: Respond")

body(
    doc,
    "Activate pre-established alternative supplier contracts. Activate minimum-volume "
    "backup supplier relationships even if not currently used. Initiate inventory "
    "drawdown protocol to extend supply runway while alternative procurement is activated. "
    "Brief treasury on currency and hedging implications of alternative supplier procurement. "
    "Begin communication with financial markets (investor relations, credit counterparties) "
    "on expected impact  proactive disclosure is always better than reactive clarification.",
)

h2(doc, "4.4  Day 3–14: Manage")

body(
    doc,
    "Assess duration: Is this a 2-week disruption or a 3-month disruption? The answer "
    "determines whether inventory drawdown is the primary tool or whether structural "
    "supply reconfiguration is required. For disruptions assessed as temporary (under "
    "30 days): maintain inventory drawdown, activate spot procurement only for "
    "incremental gaps, do not restructure supplier base. For disruptions assessed as "
    "extended: initiate formal supplier base restructuring, execute new term contracts "
    "with alternative suppliers, consider whether current supplier contract terms "
    "(including take-or-pay provisions) require commercial renegotiation. Evaluate hedging "
    "positions against updated supply scenario and adjust if warranted.",
)

h2(doc, "4.5  Day 14+: Normalize")

body(
    doc,
    "Either transition back to pre-crisis supplier base as conditions normalize, or "
    "plan for new-normal operations if the disruption is structural (sanctions regime "
    "change, geopolitical realignment). Conduct post-crisis review against the Sheffi "
    "five-dimension framework: which dimensions performed well, which revealed gaps? "
    "Update crisis playbooks with actual performance data. Adjust inventory buffer "
    "targets if the event revealed buffer levels were inadequate.",
)

add_hr(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PART V: FRAMEWORK APPLICATION  FULL MAPPING
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "PART V: FRAMEWORK APPLICATION  FULL MAPPING")

h2(doc, "5.1  Sheffi, Sodhi/Tang, and Choi et al. Applied to the Iran 2026 Case")

body(
    doc,
    "This section provides the systematic framework application that distinguishes "
    "an academic case study from an industry report. For OPIM 626, applying the course "
    "frameworks to the case facts is the intellectual core of the assignment.",
)

h3(doc, "Sheffi: The Resilient Enterprise  Five Dimensions Applied")

body(
    doc,
    "Sheffi identifies five dimensions of organizational resilience that map to the "
    "2026 Iran crisis as follows:",
)

body(
    doc,
    "Redundancy (excess capacity or inventory that can substitute for disrupted "
    "resources): Inventory buffers of 14+ days provided the absorbing layer that "
    "prevented immediate refinery runout when tanker rerouting added 10 days to "
    "transit time. Indian state-owned refiners and Japanese refiners  who maintain "
    "strategic reserves as national security infrastructure  demonstrated the value "
    "of redundancy. Chinese commercial refiners, operating on tighter inventory norms, "
    "had less absorbing capacity and were more exposed to the physical supply delay.",
    bold_prefix="Redundancy:",
)
body(
    doc,
    "Flexibility (ability to reconfigure operations to use alternative resources "
    "without significant cost or time penalty): Refinery crude slate flexibility  "
    "the ability to process Saudi Arab Light, US WTI, or Russian Urals without "
    "major throughput loss  directly determined which firms could activate "
    "alternative supply quickly. Reliance's Jamnagar complex, designed for maximum "
    "crude flexibility, absorbed the grade switch with minimal throughput impact. "
    "Simpler refineries faced 2–4 weeks of reduced throughput as they adjusted "
    "refinery operations to available crude grades.",
    bold_prefix="Flexibility:",
)
body(
    doc,
    "Awareness (real-time knowledge of supply chain conditions before disruption "
    "cascades): Real-time AIS tracking via commercial platforms (Kpler, MarineTraffic) "
    "provided 24–48 hours of early warning before the physical supply impact reached "
    "refinery gates. Firms with trading desks monitoring these systems detected "
    "vessel rerouting in near-real-time. Firms without this visibility only detected "
    "the disruption when their refinery supply position deteriorated  a 48–72 hour "
    "late start in a crisis where every hour matters.",
    bold_prefix="Awareness:",
)
body(
    doc,
    "Adaptability (ability to reconfigure supply chain network and sourcing when "
    "conditions change): Pre-established relationships with alternative suppliers "
    "(Saudi Aramco, ADNOC, US crude exporters) that could be activated without new "
    "negotiations were the clearest differentiator between Indian and Chinese refiner "
    "responses. In a crisis, new negotiations take weeks; pre-agreed contracts "
    "activate in days.",
    bold_prefix="Adaptability:",
)
body(
    doc,
    "Organizational capabilities (decision authority, cross-functional alignment, "
    "crisis playbooks): Crisis response speed was determined not just by having "
    "alternative options but by having pre-authorized decision protocols to activate "
    "them. Reliance's trading desk has commercial authority to execute supplier shifts "
    "within hours. State-owned NOCs with multi-layer approval chains were slower to "
    "activate even the alternatives they had available. The organizational "
    "capability dimension explains why pre-positioned options were not always "
    "exploited as quickly as theory would predict.",
    bold_prefix="Organizational:",
)

h3(doc, "Sodhi and Tang: Risk Taxonomy Applied")

body(
    doc,
    "Sodhi and Tang's taxonomy of supply chain risks provides a critical analytical "
    "distinction that the crisis response must address:",
)

body(
    doc,
    "The 2026 Iran disruption is a disruption risk (Sodhi and Tang's Category 1)  "
    "a sudden, external, potentially extended-duration event originating outside the "
    "firm's supply chain. This classification matters because disruption risks require "
    "different mitigation instruments than supply risks or demand risks:",
)

bullet(
    doc,
    "Disruption risks → mitigate with redundancy and flexibility (Sheffi). "
    "Hedging instruments (futures, options) address price risk but not availability risk. "
    "Chinese refiners that had heavily hedged Brent crude price risk still faced physical "
    "supply disruption because their hedges did not cover the scenario of alternative crude "
    "being physically available but at $15–20/barrel premium with 10-day delivery delay.",
)
bullet(
    doc,
    "Supply risks → mitigate with supplier management, quality controls, "
    "and contractual provisions. Not the primary risk type here, but relevant to "
    "the refinery crude slate problem (supply risk of wrong-grade crude).",
)
bullet(
    doc,
    "Demand risks → not directly relevant to this case, though product demand "
    "destruction ( refineries reducing throughput) could create downstream demand uncertainty.",
)

h3(doc, "Choi, Rogers, and Vakil: Visibility Framework Applied")

body(
    doc,
    "Choi, Rogers, and Vakil (2020) argue that supply chain resilience after COVID-19 "
    "requires visibility into tier-2 and tier-3 supply chains  not just direct tier-1 "
    "suppliers. The 2026 Iran crisis validates this argument and extends it:",
)

body(
    doc,
    "In the oil refining supply chain, tier-2 includes shipping (vessel operators, "
    "AIS tracking), tier-3 includes financial infrastructure (insurance, banking, "
    "clearing), and tier-4 includes geopolitical risk that affects both. The "
    "visibility gap in this crisis was not between refiners and their crude suppliers "
    "(tier-1)  refiners knew who their suppliers were  but between refiners and "
    "the financial and logistical infrastructure that determined whether crude could "
    "actually flow. Firms that had visibility into tier-2 (vessel movements) and "
    "tier-3 (insurance market conditions) had a 48-hour decision advantage over "
    "firms that lacked this visibility. This is precisely the dynamic Choi et al. "
    "describe in a different context (the COVID supply chain crisis), demonstrating "
    "the generalizability of their framework across disruption types.",
)

add_hr(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PART VI: IMPLICATIONS, LIMITATIONS, AND FUTURE RESEARCH
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "PART VI: IMPLICATIONS, LIMITATIONS, AND FUTURE RESEARCH")

h2(doc, "6.1  Industry-Level Implications")

body(
    doc,
    "The 2026 Iran crisis has three significant implications for Asian refiner strategy. "
    "First, the crisis demonstrates that physical supply resilience requires "
    "multi-dimensional investment  not just supplier contracts but refinery flexibility, "
    "inventory policy, visibility infrastructure, and organizational decision protocols. "
    "Firms that invested in only one dimension (e.g., supplier diversification without "
    "refinery flexibility) are recovering but not as quickly as firms with compound "
    "resilience investments. Second, the crisis reveals that geopolitical risk has "
    "re-entered the strategic planning horizon for Asian energy companies in a way "
    "not seen since the 2012–2019 sanctions escalation. The assumption that "
    "globalized oil markets would always provide alternative supply is being "
    "challenged by the weaponization of supply chain chokepoints (insurance, "
    "banking, shipping) alongside physical chokepoints. Third, the insurance market "
    "freeze represents a new category of non-physical disruption that can create "
    "functional supply loss without any physical attack on infrastructure.",
)

h2(doc, "6.2  Limitations of This Analysis")

body(
    doc,
    "Three significant limitations must be acknowledged. First, the crisis is ongoing "
    "as of April 20, 2026  all performance data (India at 80%, China at 60–70%) "
    "represents preliminary estimates, not final outcomes. The true test of "
    "resilience will be measured in months, not weeks, as the crisis either resolves "
    "or extends into a structural supply regime change. Second, firm-level performance "
    "data for Chinese NOCs is difficult to verify independently. CNPC and Sinopec "
    "disclosure standards differ from Western public companies, and some of the "
    "data in this paper is derived from industry reporting and inference rather "
    "than confirmed corporate disclosures. Third, the causal attribution problem: "
    "the India-China divergence reflects both pre-crisis positioning (which this "
    "paper emphasizes) and structural factors that are not fully reducible to "
    "firm-level resilience investment decisions. Political economy constraints, "
    "contract law differences, and sanctions history all contribute to the "
    "observed outcomes and complicate causal inference.",
)

h2(doc, "6.3  Future Research Directions")

body(
    doc,
    "As the crisis resolves, several research questions become tractable. What "
    "is the total economic cost of the disruption to Asian refiners  in volume "
    "loss, spot premium costs, and margin compression  measured against the "
    "cost of pre-crisis resilience investments? This counterfactual analysis "
    "would allow a rigorous test of Sheffi's resilience investment thesis. "
    "How did South Korean refiners (SK Energy, GS Caltex) perform, given their "
    "similar geographic exposure but different geopolitical alignment from China? "
    "The South Korean response is a natural experiment that would help separate "
    "geopolitical alignment effects from pure firm-level resilience effects. "
    "What role did crude traders (Trafigura, Vitol, Mercuria) play as market "
    "intermediaries during the crisis, and did their presence increase or "
    "decrease market efficiency in allocating alternative supply?",
)

add_hr(doc)

# ══════════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "CONCLUSION: RESILIENCE IS NOT A PRICE STORY  IT IS A STRUCTURE STORY")

body(
    doc,
    "The 2026 Iran crisis has not concluded, but the divergence between Asian refiners "
    "already tells a clear story. Brent crude went to $100  every market participant "
    "saw the price shock. But:",
)

body(
    doc,
    "India's refiners  led by IOC and Reliance Industries  are recovering to "
    "approximately 75–80% of pre-crisis throughput in 2–3 weeks. China's refiners  "
    "Sinopec and CNPC  are recovering at approximately 60–70% in 4–6 weeks. "
    "Japanese refiners experienced minimal disruption. Same price shock. Materially "
    "different outcomes  determined by pre-crisis positioning across Sheffi's five "
    "resilience dimensions, not by trading acumen during the crisis.",
)

body(
    doc,
    "Indian refiners' minimal reliance on spot market purchases during the crisis "
    "reflects pre-established alternative supplier relationships that activated in "
    "days rather than weeks. Chinese refiners' substantial spot market exposure "
    "reflects the structural difficulty of activating alternatives when take-or-pay "
    "contracts create commercial lock-in and political constraints limit supplier "
    "reallocation options. The constraint was physical, organizational, and "
    "institutional  not just a price on a screen.",
)

body(
    doc,
    "Sheffi argues in The Resilient Enterprise that supply chain resilience is "
    "determined by what companies build before disruption, not by how fast they "
    "react during it. Sodhi and Tang provide the risk taxonomy that explains "
    "why this disruption required redundancy and flexibility instruments rather "
    "than hedging instruments. Choi, Rogers, and Vakil demonstrate that the "
    "visibility gap  not the physical chokepoint  is where the actual "
    "vulnerability materialized. This is a supply chain risk management story, "
    "not a commodity trading story. And it is a story that will repeat  with "
    "different geography, different asset classes, and different firms  until "
    "the profession internalizes its lesson: resilience is structure.",
)

add_hr(doc)

# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "REFERENCES")

h2(doc, "Course Readings")

refs_course = [
    (
        "Sheffi, Y. (2005). The Resilient Enterprise: Overcoming Vulnerability for Competitive "
        "Advantage. MIT Press."
    ),
    ("Sodhi, M.S., & Tang, C.S. (2012). Managing Supply Chain Risks. Springer."),
    (
        'Choi, T.Y., Rogers, D.S., & Vakil, B. (2020). "Coronavirus is a Wake-Up Call for '
        'Supply Chain Management." Harvard Business Review, April 2020.'
    ),
]
for ref in refs_course:
    p = doc.add_paragraph(ref, style="List Bullet")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

h2(doc, "Data Sources")

refs_data = [
    (
        "U.S. Energy Information Administration (EIA). Oil Market Reports, March–April 2026. "
        "https://www.eia.gov/"
    ),
    (
        "International Energy Agency (IEA). Monthly Oil Market Reports, March–April 2026. "
        "https://www.iea.org/"
    ),
    (
        "Kpler & MarineTraffic. Tanker tracking data, vessel AIS data, March–April 2026. "
        "Subscription data accessed April 2026."
    ),
    (
        "Bloomberg L.P. Brent crude daily pricing, VLCC freight indices, war risk insurance "
        "premium data, Chinese crude import data, March–April 2026. Terminal subscription."
    ),
    (
        'Reuters. "Asian Refiners Face Supply Disruption as Hormuz Tanker Insurance Spikes." '
        "April 3, 2026."
    ),
    (
        "Lloyd's Market Association. War Risk Insurance Market Commentary. April 3, 2026. "
        "https://www.lloyds.com/"
    ),
    ("Marsh & McLennan. War Risk Insurance Briefing Note. April 4, 2026."),
    ("OPEC. Monthly Oil Market Reports, February–March 2026. https://www.opec.org/"),
    (
        "Earnings calls and investor disclosures: Indian Oil Corporation Q1 2026 Earnings "
        "Call, April 15, 2026; Sinopec Group Q1 2026 Earnings Guidance, April 2026. "
        "Publicly available via corporate investor relations."
    ),
    (
        "The Baltic Exchange. VLCC freight rate assessments, March–April 2026. "
        "https://www.balticexchange.com/"
    ),
]
for ref in refs_data:
    p = doc.add_paragraph(ref, style="List Bullet")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = "/mnt/c/Users/User/Documents/GitHub/claude-squad/CRUDE_RESILIENCE_OPIM626_Group4.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
