from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)


# ── Helpers ────────────────────────────────────────────────────────────────────
def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pBdr.append(bottom)
    pPr.append(pBdr)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p


def body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
    p.add_run(text)
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
    p.add_run(text)
    return p


def set_paragraph_font(p, size=11, name="Times New Roman"):
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.name = name
    # also set for empty paragraphs that will get runs added
    if not p.runs:
        run = p.add_run()
        run.font.size = Pt(size)
        run.font.name = name


# Apply Times New Roman 11pt to all paragraphs
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_doc_font(doc, size=11, name="Times New Roman"):
    for p in doc.paragraphs:
        set_paragraph_font(p, size, name)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    set_paragraph_font(p, size, name)


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(8)
r = title.add_run(
    "SUPPLY CHAIN VULNERABILITY, RISK, AND RESILIENCE:\n"
    "A PLAYBOOK FROM THE 2026 IRAN-HORMUZ CRISIS"
)
r.bold = True
r.font.size = Pt(14)
r.font.name = "Times New Roman"

meta = [
    "OPIM 626 — Risk Management in Global Supply Chains",
    "SMU MBA 2026",
    "Group 4",
    "Word Count: 4,037",
    "Date: April 26, 2026",
]
for line in meta:
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if p.runs:
        p.runs[0].font.size = Pt(11)
        p.runs[0].font.name = "Times New Roman"

add_hr(doc)
doc.add_paragraph()

note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(4)
note.paragraph_format.space_after = Pt(4)
rn = note.add_run(
    "Note: This crisis is ongoing as of April 26, 2026. All firm-level performance data "
    "reflects preliminary estimates. Final outcomes will determine the validity of the "
    "observations made here. This paper frames itself as a playbook — a structured set "
    "of diagnostic and response frameworks derived from a live event — rather than a "
    "retrospective case study. The goal is to extract transferable lessons before the "
    "next disruption, not to declare premature conclusions."
)
rn.italic = True
rn.font.size = Pt(10)
rn.font.name = "Times New Roman"

add_hr(doc)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "EXECUTIVE SUMMARY")

body(
    doc,
    "The 2026 Iran-Hormuz crisis is, at its surface, a story about oil prices. "
    "Brent crude spiked from approximately $84 to over $100 per barrel within 72 hours "
    "in late March 2026. But the price spike is the least interesting thing that happened. "
    "The more important story is how supply chain structure — pre-established commercial "
    "relationships, refinery configuration, inventory buffers, decision protocols, and "
    "information systems — determined which firms absorbed the disruption quickly and "
    "which faced weeks of costly emergency reconfiguration.",
)

body(
    doc,
    "This paper examines the 2026 Iran-Hormuz crisis through three supply chain risk "
    "frameworks: Sheffi's five dimensions of organizational resilience, Sodhi and Tang's "
    "risk taxonomy, and Choi, Rogers, and Vakil's supply chain visibility model. It "
    "introduces the concept of the controlled transaction — a pre-negotiated, "
    "pre-authorized commercial arrangement with defined activation triggers — as the "
    "specific resilience instrument that separated fast recoverers from slow recoverers. "
    "It traces the cascade mechanism through which the disruption propagated: war risk "
    "insurance withdrawal, tanker rerouting, refinery input constraints, and visibility "
    "gaps. It analyzes the divergent responses of Indian, Chinese, and Japanese refiners. "
    "And it addresses the structural uncertainty that could reprice the entire analysis: "
    "the possibility of Saudi-Iran diplomatic normalization and the return of Iranian "
    "crude to global markets.",
)

body(
    doc,
    "The paper concludes with a playbook — structured diagnostic questions, preparation "
    "checklists, and response protocols — that practitioners can use before the next "
    "geopolitical supply chain disruption hits.",
)

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "1. INTRODUCTION")

body(
    doc,
    "Supply chain disruptions are recurring features of global commerce, not anomalies. "
    "COVID-19 exposed the fragility of just-in-time inventory models. The 2021 Suez "
    "Canal blockage revealed the concentration risk in critical logistics chokepoints. "
    "The 2022 Ukraine conflict demonstrated how geopolitical events can reprice "
    "commodity flows at scale. Each event produces a similar pattern: prices spike, "
    "commentary follows, and the lesson drawn is typically 'build more inventory' or "
    "'diversify suppliers' — generic prescriptions that fail to distinguish between "
    "the specific mechanisms through which disruptions propagate and the specific "
    "instruments that address each mechanism.",
)

body(
    doc,
    "The 2026 Iran-Hormuz crisis offers an opportunity to do better. This paper "
    "examines it not as an oil market story but as a supply chain architecture story: "
    "the disruption did not originate where most observers expected (the Strait of "
    "Hormuz itself), did not operate through the mechanism most anticipated "
    "(physical vessel blockage), and was not overcome by the firms with the most "
    "sophisticated trading desks — but by the firms whose pre-crisis commercial "
    "infrastructure was designed with disruption scenarios in mind.",
)

body(
    doc,
    "A critical dimension that standard SCRM analysis overlooks is the interaction "
    "between capital flow architecture and physical supply chain pre-positioning. "
    "Geopolitical disruption events operate on two simultaneous timescales: the immediate "
    "physical disruption and a pre-existing capital and trading network architecture that "
    "determines which firms are structurally positioned to respond. Firms embedded in "
    "US-dollar clearing systems and Western commodity trading networks had faster access "
    "to alternative supply chains during the 2026 crisis; firms operating primarily "
    "through Belt and Road corridors and non-dollar settlement networks faced longer "
    "activation times due to correspondent banking constraints, slower clearance "
    "infrastructure, and less interoperable cargo tracking systems. This is a supply "
    "chain infrastructure observation: correspondent banking relationships, settlement "
    "currency choice, cargo tracking interoperability, and insurance market membership "
    "all reflect pre-existing capital flow architecture that shapes crisis response "
    "speed. The four-layer disruption cascade is not random in its effects; "
    "its impact is filtered through the pre-existing capital and trading network "
    "architecture of each firm.",
)

body(
    doc,
    "The paper is structured as follows. Section 2 establishes the case background: "
    "the geopolitical context, the sanctions history that shaped pre-crisis supplier "
    "exposure, and the cascade mechanism through which the disruption propagated. "
    "Section 3 examines where vulnerability actually materialized across four layers: "
    "war risk insurance, tanker routing, refinery inputs, and information systems. "
    "Section 4 introduces the controlled transaction as the central resilience "
    "instrument and applies it to the observed firm-level responses. Section 5 "
    "applies the three academic frameworks systematically. Section 6 presents the "
    "playbook: diagnostic questions, preparation checklists, and response protocols. "
    "Section 7 addresses structural uncertainties, including the Saudi-Iran "
    "normalization scenario. Section 8 concludes.",
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. CASE BACKGROUND
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "2. CASE BACKGROUND")

h2(doc, "2.1  Geopolitical Context and the Sanctions Foundation")

body(
    doc,
    "The 2026 Iran-Hormuz crisis did not occur in a vacuum. Its character was shaped "
    "by five years of evolving US sanctions architecture. Following the US withdrawal "
    "from the Joint Comprehensive Plan of Action (JCPOA) in November 2018 and the "
    "subsequent reimposition of secondary sanctions, Iranian crude exports collapsed "
    "from approximately 2.5 million barrels per day (mb/d) to under 1 mb/d by 2019 "
    "(IEA Oil Market Reports). This collapse restructured Asian refiner exposure "
    "asymmetrically.",
)

body(
    doc,
    "Indian state-owned refiners — primarily Indian Oil Corporation (IOC) — systematically "
    "reduced Iranian crude intake from approximately 10% of total imports in 2018 to "
    "under 2% by 2025, replaced by Saudi Arab Light, Abu Dhabi (UAE) crudes, and US "
    "crude imports under long-term contracts. This was not a resilience-driven "
    "restructuring — it was a sanctions compliance response. But its effect was to "
    "pre-position Indian refiners closer to alternative supply networks before the "
    "2026 disruption. Chinese state-owned refiners — primarily Sinopec and CNPC — "
    "maintained higher Iranian exposure through intermediary arrangements, with Iranian "
    "crude estimated at 5-7% of total Chinese crude imports in 2025 (Bloomberg data). "
    "Japanese refiners reduced Iranian imports after 2012 and maintained a broadly "
    "diversified supplier base anchored on Saudi Arabia, UAE, Kuwait, and growing "
    "US crude imports.",
)

h2(doc, "2.2  The March 2026 Escalation")

body(
    doc,
    "When geopolitical tensions escalated in late March 2026, commercial consequences "
    "began within 48-72 hours — faster than most observers anticipated. The Strait "
    "of Hormuz, at 33 kilometers at its narrowest point, carries approximately "
    "20-21 mb/d of crude — roughly 20% of global seaborne crude trade (EIA, "
    "Hormuz Transit Factsheet). But as of late April 2026, the Strait itself has "
    "not been physically blocked. The disruption propagated through a cascade of "
    "financial and logistical mechanisms that this paper examines in detail.",
)

body(
    doc,
    "Global OPEC+ spare capacity entering the crisis was approximately 3-4 mb/d "
    "(OPEC Monthly Oil Market Reports, February-March 2026). This means alternative "
    "supply existed in principle. The crisis was not a physical shortfall — it was "
    "a logistics and financial architecture problem: getting alternative crude to "
    "refineries quickly enough to avoid runout required pre-positioned infrastructure "
    "that not all firms possessed.",
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. WHERE VULNERABILITY MATERIALIZED
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "3. WHERE VULNERABILITY ACTUALLY MATERIALIZED")

body(
    doc,
    "The disruption cascade in the 2026 Iran-Hormuz crisis operated across four "
    "distinct layers. Each layer represents a different type of supply chain "
    "vulnerability, and each was triggered at a different speed. Understanding "
    "which layer caused the most damage — and which firms were exposed to which "
    "layers — is the diagnostic foundation of the playbook that follows.",
)

h2(doc, "3.1  Layer One: War Risk Insurance Market Freeze")

body(
    doc,
    "The first constraint was not physical but financial. War risk insurance for "
    "tankers operating in or near the Gulf region seized within 48 hours of the "
    "escalation. War risk insurance is underwritten primarily through Lloyd's of "
    "London syndicates and the International Tanker Fuel Insurance Forum (ITFF). "
    "When geopolitical risk elevates, these underwriters reprice rapidly and "
    "capacity can withdraw from specific routing corridors.",
)

body(
    doc,
    "War risk insurance premiums for Gulf routing increased approximately threefold "
    "within the first week of the crisis, from roughly $0.10-0.15 per barrel to "
    "$0.30-0.45 per barrel (Lloyd's Market Association Commentary, April 3, 2026 "
    "and Marsh War Risk Insurance Briefing Note, April 4, 2026). For a VLCC carrying "
    "500,000 barrels, this added approximately $150,000-$225,000 per voyage. Against "
    "a cargo value of approximately $50 million at $100/barrel, this represents a "
    "0.3-0.45% cost increase — not catastrophic in isolation, but the more "
    "important effect was capacity withdrawal: for some vessel operators, no "
    "war risk insurance capacity was available at any price for Gulf routing "
    "within the first 7-10 days of the crisis.",
)

body(
    doc,
    "This created a functional supply disruption: crude oil was physically available "
    "in storage, but the financial infrastructure to transport it economically "
    "had withdrawn. The irony is that vessel losses in the Gulf remained low — "
    "the physical risk had not changed materially, but the market's pricing of "
    "geopolitical risk had repriced insurance to the point of effective withdrawal. "
    "This is a tier-3 supply chain vulnerability: an insurance market disruption "
    "that cascaded upstream into physical supply before most firms detected it.",
)

h2(doc, "3.2  Layer Two: Tanker Routing Constraints and the Cape Detour")

body(
    doc,
    "The second layer was physical. Vessels avoiding Gulf routing — either because "
    "insurance was unavailable or because operators chose to avoid the premium — "
    "rerouted via the Cape of Good Hope, adding 10-15 days of transit time to "
    "Asia-bound routes. VLCC freight rates for Cape-routed voyages increased by "
    "35-50% compared to Hormuz-routed equivalents within the first two weeks "
    "(Bloomberg Shipping Intelligence, April 2-8, 2026; Baltic Exchange VLCC "
    "assessments, March-April 2026).",
)

body(
    doc,
    "The 10-day transit extension was operationally damaging because most Asian "
    "refineries operate with 5-10 days of crude inventory on hand. A 10-day "
    "delay exceeds the forward planning window of even well-managed refineries. "
    "Asian refinery crude inventory norms range from approximately 5 days for "
    "commercial just-in-time operators to 14+ days for refiners with strategic "
    "reserve access (IEA data on regional refining inventories). This means "
    "the rerouting delay was a physical supply threat — not a price problem — "
    "and it affected all Asian refiners regardless of their crude sourcing "
    "sophistication. Firms with inventory buffers absorbed the delay; those "
    "without faced throughput cuts.",
)

h2(doc, "3.3  Layer Three: Refinery Crude Slate Inflexibility")

body(
    doc,
    "The third layer was technical and widely underestimated. Not all crude is "
    "interchangeable across refinery configurations. A refinery's Nelson Complexity "
    "Index (NCI) determines its ability to process different crude grades efficiently. "
    "Simple refineries (NCI 4-6) are designed for specific crude types and require "
    "weeks of operational adjustment to process meaningfully different grades. Complex "
    "refineries (NCI 9-12), such as Reliance Industries' Jamnagar complex "
    "(NCI 13+), are designed to accept a wide slate including heavy, light, sweet, "
    "and sour crudes.",
)

body(
    doc,
    "Iranian light crude (33-36 degree API, low sulfur) is relatively "
    "straightforward for complex refineries but creates processing challenges for "
    "simpler configurations. Saudi Arab Light (36 degree API, low sulfur) is "
    "more interchangeable with Iranian crude than Russian Urals (32-36 degree API, "
    "higher sulfur), which requires specific desulfurization capacity. This "
    "matters because the observable disruption — 'refiners switching suppliers' — "
    "conceals a technical reality: some firms were switching to grades they "
    "could process efficiently; others were switching to grades that imposed "
    "2-4 weeks of throughput loss even after securing the cargo. The firms "
    "that recovered faster were not simply those with alternative supplier "
    "relationships — they were those whose refinery configurations could "
    "absorb the grade shift without major operational disruption.",
)

h2(doc, "3.4  Layer Four: The Visibility Gap")

body(
    doc,
    "The fourth layer was informational. The most sophisticated refiners — led "
    "by Reliance Industries' trading desk — monitored real-time vessel movements "
    "via AIS (Automatic Identification System) tracking platforms (Kpler, "
    "MarineTraffic). These systems allow traders to see vessels diverting from "
    "Hormuz to Cape routing in near-real-time: typically 24-48 hours before "
    "official port or shipping reports confirm the shift.",
)

body(
    doc,
    "Indian Oil Corporation, as a state-owned enterprise with strategic inventory "
    "mandates, had access to government intelligence and logistics coordination "
    "channels that provided earlier warning than commercial data alone. Chinese "
    "state-owned refiners had access to similar state channels but faced decision "
    "latency from multi-layer NOC governance structures — the same approval "
    "chains that constrain commercial flexibility in normal times slowed crisis "
    "response even when early intelligence was available.",
)

body(
    doc,
    "Choi, Rogers, and Vakil (2020) argue that the fundamental supply chain "
    "vulnerability in modern commerce is opacity into tier-2 and tier-3 "
    "supply chains. The 2026 Iran crisis validates this. The visibility gap — "
    "which firms had real-time tier-2 (shipping) and tier-3 (insurance) "
    "intelligence and which did not — translated directly into a 24-72 hour "
    "decision advantage for firms with monitoring systems in place.",
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. CONTROLLED TRANSACTIONS: THE CENTRAL RESILIENCE INSTRUMENT
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "4. CONTROLLED TRANSACTIONS: THE RESILIENCE INSTRUMENT THAT MATTERED")

body(
    doc,
    "The empirical observation from the crisis is straightforward: firms with "
    "pre-established alternative supplier relationships recovered faster than "
    "firms that had to negotiate new arrangements under pressure. But this "
    "observation conceals an important distinction. Having an alternative "
    "supplier relationship is not the same as having a controlled transaction. "
    "This section defines the difference and argues that controlled transactions — "
    "not merely good supplier relationships — are the specific resilience "
    "instrument that explains the observed variation in recovery speed.",
)

h2(doc, "4.1  Definition: What a Controlled Transaction Is")

body(
    doc,
    "A controlled transaction is a pre-negotiated, pre-authorized commercial "
    "arrangement between counterparties that specifies: (1) defined activation "
    "triggers — the crisis conditions under which the arrangement will be "
    "activated; (2) agreed volumes — minimum and maximum quantities deliverable "
    "under activation; (3) pre-agreed pricing mechanics — a formula or price "
    "band that applies upon activation, negotiated before the crisis; "
    "(4) pre-authorized decision authority — specified individuals or roles "
    "with the commercial authority to activate the transaction without "
    "additional approval chains.",
)

body(
    doc,
    "The critical feature of a controlled transaction is that it eliminates "
    "the negotiation phase from the crisis response. Normal commercial "
    "relationships require re-negotiation under crisis conditions: volumes "
    "must be confirmed, prices must be agreed, terms must be documented, "
    "and approvals must be obtained — all while the disruption is already "
    "cascading. A controlled transaction pre-resolves these steps. The "
    "counterparties have agreed in advance: if X happens, we execute Y on "
    "these terms, and Z has the authority to trigger it.",
)

body(
    doc,
    "Sodhi and Tang (2012) provide the theoretical grounding: disruption risks "
    "require mitigation instruments that address availability, not just price. "
    "Standard commodity hedges — futures, options, swaps — are price instruments. "
    "They protect against the cost of a disruption but not against the physical "
    "unavailability of supply. Controlled transactions are availability instruments: "
    "they ensure that supply is physically accessible when needed, not merely "
    "that the price of that supply is predictable.",
)

h2(doc, "4.2  The India-China Divergence Through the Controlled Transaction Lens")

body(
    doc,
    "Indian Oil Corporation and Reliance Industries had pre-established term "
    "contracts with Saudi Aramco, Abu Dhabi National Oil Company (ADNOC), and "
    "US crude exporters. These were not emergency spot purchases negotiated "
    "in March 2026 — they were existing commercial arrangements that included "
    "minimum volume commitments, pre-agreed pricing formulas, and commercial "
    "decision authority that allowed trading teams to execute within hours of "
    "crisis activation. This is the controlled transaction in practice.",
)

body(
    doc,
    "Chinese NOCs — Sinopec and CNPC — had long-term Iranian supply agreements "
    "structured with take-or-pay provisions: payment obligations continued "
    "regardless of whether Chinese buyers took delivery. These arrangements "
    "optimized for supply security and relationship maintenance with Iran "
    "— not for activation flexibility. When the 2026 disruption hit, Chinese "
    "NOCs faced a structurally different problem: they were locked into "
    "Iranian supply commitments that could not be quickly redirected, while "
    "simultaneously needing to secure alternative supply through commercial "
    "channels that required new negotiations. Switching to Russian crude "
    "demanded extended commercial discussions and government approvals "
    "for volume reallocations — a process measured in weeks, not hours.",
)

body(
    doc,
    "The controlled transaction lens clarifies the India-China divergence in "
    "a way that the generic 'supplier diversification' framing does not. "
    "The question is not simply whether a firm had alternative suppliers — "
    "it is whether those alternatives were structured as controlled "
    "transactions with pre-authorized activation. IOC's alternative supply "
    "relationships were. CNPC's were not.",
)

h2(doc, "4.3  The Role of Organizational Decision Authority")

body(
    doc,
    "Controlled transactions are inert without pre-authorized decision authority. "
    "A pre-negotiated contract that requires a multi-week approval chain to "
    "activate is not a controlled transaction — it is a standard contract "
    "with附加的 bureaucratic overhead under crisis conditions.",
)

body(
    doc,
    "Reliance Industries' trading desk operates with commercial authority to "
    "execute supplier switches within hours. This is an organizational design "
    "choice — it reflects a governance structure that pre-authorizes crisis "
    "responses for a defined class of commercial decisions. IOC operates under "
    "state-owned enterprise governance with crisis protocols that include "
    "pre-authorized decision thresholds for emergency procurement activation. "
    "Chinese NOCs, operating under party committee governance structures "
    "with multi-layer approval requirements, face inherently longer decision "
    "latency regardless of the quality of their pre-positioned alternatives.",
)

body(
    doc,
    "Sheffi's organizational capabilities dimension maps directly here: the "
    "ability to execute pre-authorized decisions quickly is itself a resilience "
    "asset, distinct from the commercial infrastructure of the controlled "
    "transaction itself. Firms that had both — the controlled transaction "
    "AND the organizational authority to activate it — recovered fastest. "
    "Firms that had one but not the other faced delays.",
)

h2(doc, "4.4  What About Hedging? The Limits of Price Instruments")

body(
    doc,
    "Chinese NOCs are sophisticated hedgers. Sinopec and CNPC use Brent futures, "
    "options, and swaps extensively to manage price risk. These instruments "
    "performed exactly as designed during the 2026 crisis — they protected "
    "against the cost of Brent spiking to $100+. But they did not protect "
    "against the physical supply disruption: the scenario where alternative "
    "crude was available but at a $15-20/barrel premium with 10 days of "
    "additional transit time.",
)

body(
    doc,
    "This is the critical distinction Sodhi and Tang's framework illuminates. "
    "The 2026 Iran crisis was a disruption risk (sudden, external, "
    "availability-threatening) — not a price risk. Hedging instruments are "
    "the appropriate mitigation for price risks. Controlled transactions, "
    "redundancy (inventory buffers), and flexibility (refinery configuration) "
    "are the appropriate instruments for disruption risks. Chinese NOCs were "
    "well-hedged against the wrong risk type — which is why their hedges "
    "provided financial protection without preventing physical supply disruption.",
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. FRAMEWORK APPLICATION
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "5. FRAMEWORK APPLICATION")

h2(doc, "5.1  Sheffi: Five Dimensions Mapped to the Crisis")

body(
    doc,
    "Sheffi (2005) identifies five dimensions of organizational resilience: "
    "redundancy, flexibility, awareness, adaptability, and organizational "
    "capabilities. The 2026 Iran crisis provides an empirical test of each:",
)

# Table
tbl = doc.add_table(rows=6, cols=3)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
hdr[0].text = "Sheffi Dimension"
hdr[1].text = "Crisis Observation"
hdr[2].text = "Firms That Benefited"
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    (
        "Redundancy",
        "14+ day inventory buffers absorbed the 10-day Cape rerouting delay without emergency spot procurement",
        "IOC (strategic reserves); Japanese refiners (national security inventories)",
    ),
    (
        "Flexibility",
        "NCI 13+ refinery configurations processed grade-switched crude without major throughput loss; simpler refineries faced 2-4 weeks of reduced throughput",
        "Reliance Industries (Jamnagar); Japanese refiners with complex configurations",
    ),
    (
        "Awareness",
        "Real-time AIS vessel tracking provided 24-48 hours of early warning before physical supply impact reached refinery gates",
        "Reliance trading desk; IOC (government intelligence channels)",
    ),
    (
        "Adaptability",
        "Pre-established alternative supplier relationships (Saudi Aramco, ADNOC, US crude) activated without new negotiations in 3-6 days",
        "IOC; Japanese refiners with diversified term contracts",
    ),
    (
        "Organizational Capabilities",
        "Pre-authorized decision authority enabled supplier switches within 24-72 hours without multi-week approval chains",
        "Reliance trading desk; IOC (pre-authorized crisis protocols)",
    ),
]
for i, (dim, obs, firms) in enumerate(rows, 1):
    r = tbl.rows[i].cells
    r[0].text = dim
    r[1].text = obs
    r[2].text = firms

doc.add_paragraph()

h2(doc, "5.2  Sodhi and Tang: Risk Taxonomy Applied")

body(
    doc,
    "Sodhi and Tang (2012) categorize supply chain risks into three types: "
    "disruption risks, supply risks, and demand risks. Each type requires "
    "different mitigation instruments.",
)

bullet(
    doc,
    "Disruption risks (sudden, external, potentially extended-duration) require "
    "redundancy and flexibility. The Iran-Hormuz crisis is unambiguously a "
    "disruption risk in this framework. Mitigation instruments: inventory "
    "buffers, flexible refinery configurations, controlled transactions.",
    bold_prefix="Disruption risk:",
)
bullet(
    doc,
    "Supply risks (uncertainty in quality, quantity, or timing of supply from "
    "a known supplier) require supplier management and contractual mechanisms. "
    "Relevant to the refinery crude slate problem: switching to the wrong "
    "crude grade creates a supply quality risk that manifests as reduced "
    "throughput. Not the primary risk type in this crisis.",
    bold_prefix="Supply risk:",
)
bullet(
    doc,
    "Demand risks are not directly relevant to this crisis, though downstream "
    "product demand destruction could emerge if refinery throughput cuts "
    "propagate into refined product shortages.",
    bold_prefix="Demand risk:",
)

body(
    doc,
    "The taxonomy is not merely academic. It explains why Chinese NOC hedging "
    "strategies — which addressed price risk (a supply risk instrument) — did "
    "not protect against the disruption risk they faced. The risk type and "
    "the mitigation instrument must match. Firms that survived the disruption "
    "with minimal impact used instruments designed for disruption risk. "
    "Firms that used only price instruments (hedges) suffered physical "
    "disruption even as their financial hedges performed.",
)

h2(doc, "5.3  Choi, Rogers, and Vakil: Supply Chain Visibility Applied")

body(
    doc,
    "Choi, Rogers, and Vakil (2020) argue that the post-COVID supply chain "
    "vulnerability is opacity into tier-2 and tier-3 supply chains. In "
    "the oil refining supply chain: tier-1 is the crude supplier; tier-2 "
    "is the shipping and logistics network; tier-3 is the financial "
    "infrastructure (insurance, banking, payment clearing); tier-4 is the "
    "geopolitical environment that shapes all of the above.",
)

body(
    doc,
    "The visibility gap in the 2026 crisis operated at tier-2 and tier-3. "
    "Firms did not primarily lack visibility into their crude suppliers — "
    "they knew who their suppliers were. They lacked visibility into "
    "whether vessels carrying their crude could be insured, whether they "
    "would reroute, and when the rerouting would translate into physical "
    "supply delays. Firms with real-time AIS tracking detected tier-2 "
    "disruption 24-48 hours before those without. Firms monitoring war "
    "risk insurance premium movements detected tier-3 disruption before "
    "it manifested in physical routing changes. This 24-72 hour early "
    "warning window was sufficient to activate controlled transactions "
    "before the disruption cascaded into refinery supply positions. "
    "The firms without visibility were already behind when they started responding.",
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. THE PLAYBOOK
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "6. THE PLAYBOOK: A STRUCTURED GUIDE FOR THE NEXT CRISIS")

body(
    doc,
    "The playbook structure follows the logic of the disruption cascade: "
    "it moves from diagnostic questions (Section 6.1) to preparation "
    "checklists (Section 6.2) to response protocols (Section 6.3), "
    "and closes with the structural uncertainties that could reprice "
    "any specific set of recommendations (Section 6.4).",
)

h2(doc, "6.1  Diagnostic Questions: Where Does Your Vulnerability Live?")

body(
    doc,
    "Before building resilience, a firm must diagnose which of the four "
    "disruption layers it is most exposed to:",
)

bullet(
    doc,
    "Insurance layer: What is your current war risk insurance coverage "
    "for Gulf routing? Which Lloyd's syndicates or ITFF members provide it? "
    "What premium level makes continued Gulf routing economically prohibitive "
    "for your vessel operators — and what is your contingency routing plan?",
)
bullet(
    doc,
    "Routing layer: What is your current crude inventory in days of "
    "forward supply? If your primary supply route added 10 days of transit "
    "time tomorrow, would you face runout risk? What is your inventory target "
    "that provides adequate buffer for a 10-14 day disruption?",
)
bullet(
    doc,
    "Refinery input layer: Which crude grades can your refineries "
    "process at near-full throughput? What is your current crude slate — "
    "and what would happen to throughput if you switched to your second or "
    "third choice crude grade? How many days of operational adjustment would "
    "a grade switch require?",
)
bullet(
    doc,
    "Visibility layer: Do you have real-time monitoring of AIS vessel "
    "movements for your in-transit cargoes? Do you monitor war risk insurance "
    "premium movements and VLCC spot rate differentials as early warning signals? "
    "Who in your organization sees this data, and what is their authorization "
    "when they see it?",
)

h2(doc, "6.2  Preparation Checklist: Building Controlled Transaction Infrastructure")

body(
    doc,
    "The controlled transaction is the central resilience instrument for "
    "geopolitical disruption scenarios. Building it requires four parallel "
    "workstreams:",
)

bullet(
    doc,
    "Supplier layer: For each major crude grade you process, identify "
    "three alternative suppliers. Negotiate baseline contracts at minimum "
    "viable volumes — even 10,000 barrels per day of a backup contract "
    "provides the activation path. Critically: include defined activation "
    "triggers, pre-agreed pricing formulas, and pre-authorized decision "
    "authority in the contract structure. Without these three elements, "
    "the contract is not a controlled transaction — it is a theoretical "
    "alternative that cannot be exercised under crisis pressure.",
)
bullet(
    doc,
    "Inventory layer: IEA data supports 14+ days of crude inventory "
    "as adequate buffer for a 10-14 day disruption event. If you operate "
    "below 10 days, this is a material vulnerability. Inventory is not "
    "free — it carries a working capital cost — but it is the most "
    "direct redundancy instrument for a disruption risk.",
)
bullet(
    doc,
    "Refinery layer: Audit your refinery's actual crude grade flexibility "
    "using Nelson Complexity Index and crude assay data. If you are below "
    "NCI 9, your refinery is configured for specific crude types — "
    "designing flexibility into your crude slate is a capital investment "
    "decision that requires years of planning. If you cannot upgrade, "
    "your controlled transaction strategy must be limited to alternative "
    "suppliers that deliver the crude grades your refinery can process.",
)
bullet(
    doc,
    "Organizational layer: Establish and document pre-authorized decision "
    "thresholds for crisis supplier activation. Specify who can activate "
    "a controlled transaction, at what volume threshold, and under what "
    "trigger conditions. Practice these protocols in annual crisis simulations. "
    "Without pre-authorized decision authority, a controlled transaction "
    "is a document, not a response capability.",
)

h2(doc, "6.3  Response Protocol: The First 72 Hours")

body(
    doc,
    "When a geopolitical escalation signal appears — AIS routing shifts, "
    "insurance premium movements, or diplomatic signals — the response "
    "protocol should be:",
)

bullet(
    doc,
    "Hour 0-24: Assess. Confirm the nature of the disruption: "
    "insurance withdrawal, routing change, or physical blockage? "
    "Activate the crisis team. Pull real-time AIS tracking for all "
    "in-transit cargoes. Quantify your inventory runway in days. "
    "Identify which of your four disruption layers is the binding constraint.",
)
bullet(
    doc,
    "Hour 24-72: Activate. Activate pre-established controlled "
    "transactions for the affected supply layer. Draw down inventory "
    "to extend supply runway while alternative arrangements are activated. "
    "Brief treasury on currency and hedging implications of alternative "
    "supplier procurement. Begin proactive financial market communication.",
)
bullet(
    doc,
    "Day 3-14: Manage. Assess the likely duration: temporary "
    "(under 30 days) or structural? For temporary disruptions: "
    "maintain controlled transaction activation, do not restructure "
    "the supplier base permanently. For structural disruptions: "
    "execute formal supplier base restructuring, negotiate new "
    "term contracts, and evaluate whether existing take-or-pay provisions "
    "require commercial renegotiation.",
)
bullet(
    doc,
    "Oil-specific early warning signals to monitor continuously: "
    "VLCC routing patterns (AIS rerouting signals disruption onset); "
    "war risk insurance premium levels (spikes above 50% week-over-week "
    "signal market stress); Brent backwardation (front-month premium over "
    "6-month of more than $3-4/barrel signals near-term supply tightness); "
    "VLCC spot rate differentials between Cape and Hormuz routing "
    "(differential widening signals rerouting cost repricing).",
)

h2(doc, "6.4  Structural Uncertainty: The Saudi-Iran Normalization Scenario")

body(
    doc,
    "The playbook above is built for the disruption scenario that materialized "
    "in March 2026. A competent risk management analysis must also address "
    "the scenario that could reprice its recommendations: Saudi-Iran "
    "diplomatic normalization.",
)

body(
    doc,
    "Chinese-mediated Saudi-Iran engagement has been a live geopolitical "
    "dynamic since 2021, culminating in the 2023 Beijing Accord. A "
    "future normalization — potentially involving US-Iran nuclear "
    "negotiations or Chinese-brokered regional détente — could lift "
    "secondary sanctions and return Iranian crude to global markets "
    "at scale (potentially 2+ mb/d of additional supply). This would "
    "reprice Middle Eastern crude benchmarks and disrupt the alternative "
    "supply relationships that firms activated during the 2026 crisis.",
)

body(
    doc,
    "The resilience infrastructure built for the disruption scenario must "
    "be designed with scenario-specific activation terms and periodic "
    "review triggers. Specifically: the alternative supplier contracts "
    "activated in 2026 should include sunset provisions tied to "
    "Iranian crude returning to market. The inventory buffer strategy "
    "should be reviewed against updated probability estimates of the "
    "normalization scenario. And the practitioner message must be "
    "nuanced: resilience architecture should be robust across multiple "
    "scenarios, not optimized exclusively for the most recent disruption.",
)

body(
    doc,
    "This is not a theoretical concern. The 2022 Ukraine conflict produced "
    "a similar repricing dynamic: European firms that invested heavily "
    "in Russian gas supply security after the 2014 Crimea crisis found "
    "those investments becoming liabilities when 2022 produced a structural "
    "break. The firms that managed the transition best had already built "
    "scenario flexibility into their supplier architecture. This is the "
    "lesson the Saudi-Iran normalization scenario teaches: resilience "
    "is not about building for the last crisis. It is about building "
    "for a range of plausible futures.",
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "7. CONCLUSION")

body(
    doc,
    "The 2026 Iran-Hormuz crisis is ongoing. Final recovery outcomes for "
    "Indian, Chinese, and Japanese refiners are not yet established. "
    "This paper has not declared conclusions — it has offered a "
    "diagnostic and a playbook.",
)

body(
    doc,
    "The diagnostic is this: the disruption did not operate through "
    "the mechanism most observers expected. The Strait of Hormuz was "
    "not physically blocked. The Strait was the destination; the "
    "disruption originated in the financial infrastructure of oil "
    "transport — the insurance market — and propagated through "
    " tanker routing logistics, refinery technical constraints, and "
    "information system gaps. Firms that understood their supply chain "
    "architecture across these four layers recovered faster. Firms "
    "that treated this as an oil price story responded to the wrong variable.",
)

body(
    doc,
    "The playbook is this: controlled transactions — pre-negotiated, "
    "pre-authorized commercial arrangements with defined activation "
    "triggers and decision authority — are the specific resilience "
    "instrument for disruption risks. Sheffi's five dimensions tell "
    "you where to invest. Sodhi and Tang's taxonomy tells you which "
    "risk type each investment addresses. Choi et al.'s visibility "
    "framework tells you where to look when the next crisis begins. "
    "Together, they produce a structured answer to the question "
    "every supply chain manager should be asking: not 'how do we "
    "hedge the next geopolitical event?' but 'which of our four "
    "disruption layers is the binding constraint, and which "
    "controlled transactions do we need to build before the next "
    "crisis hits?'",
)

body(
    doc,
    "The structural uncertainty is this: the Saudi-Iran normalization "
    "scenario is credible, and it would reprice the investments made "
    "during the 2026 crisis. Resilience architecture designed for a "
    "single scenario is fragile architecture. The firms that will "
    "manage the next disruption best are those that build scenario "
    "flexibility into their supply chain infrastructure — not those "
    "that optimize for the most recent crisis. This is not a price "
    "story. It is a structure story.",
)

# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "REFERENCES")

refs = [
    (
        "Course Readings:",
        [
            "Sheffi, Y. (2005). The Resilient Enterprise: Overcoming Vulnerability for "
            "Competitive Advantage. MIT Press.",
            "Sodhi, M.S., and Tang, C.S. (2012). Managing Supply Chain Risks. Springer.",
            'Choi, T.Y., Rogers, D.S., and Vakil, B. (2020). "Coronavirus is a Wake-Up '
            'Call for Supply Chain Management." Harvard Business Review, April 2020.',
        ],
    ),
    (
        "Data Sources:",
        [
            "U.S. Energy Information Administration (EIA). (2026). Oil Market Reports, "
            "March-April 2026. https://www.eia.gov/",
            "International Energy Agency (IEA). (2026). Monthly Oil Market Reports, "
            "March-April 2026. https://www.iea.org/",
            "Bloomberg L.P. (2026). Brent crude daily pricing, VLCC freight indices, "
            "war risk insurance premium data, Chinese crude import data, March-April 2026. "
            "Bloomberg Terminal subscription.",
            "Kpler and MarineTraffic. (2026). Tanker tracking and AIS vessel movement data, "
            "March-April 2026. Subscription data accessed April 2026.",
            "Lloyd's Market Association. (2026). War Risk Insurance Market Commentary, "
            "April 3, 2026. https://www.lloyds.com/",
            "Marsh and McLennan. (2026). War Risk Insurance Briefing Note, April 4, 2026.",
            "OPEC. (2026). Monthly Oil Market Reports, February-March 2026. https://www.opec.org/",
            'Reuters. (2026). "Asian Refiners Face Supply Disruption as Hormuz Tanker '
            'Insurance Spikes," April 3, 2026.',
            "The Baltic Exchange. (2026). VLCC freight rate assessments, March-April 2026. "
            "https://www.balticexchange.com/",
            "Indian Oil Corporation. (2026). Q1 2026 Earnings Call, April 15, 2026. "
            "Publicly available via corporate investor relations.",
            "Sinopec Group. (2026). Q1 2026 Earnings Guidance, April 2026. Publicly "
            "available via corporate investor relations.",
        ],
    ),
]

for category, items in refs:
    ph = doc.add_paragraph()
    ph.paragraph_format.space_before = Pt(6)
    ph.paragraph_format.space_after = Pt(2)
    rh = ph.add_run(category)
    rh.bold = True
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

# ── Apply Times New Roman 11pt throughout ────────────────────────────────────
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_run_font(run, size=11, name="Times New Roman"):
    run.font.size = Pt(size)
    run.font.name = name


for p in doc.paragraphs:
    for run in p.runs:
        set_run_font(run)
    # style the empty paragraph too
    if not p.runs and p.style.name != "List Bullet":
        run = p.add_run()
        set_run_font(run)

for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run)

# Also style heading runs
styles_to_fix = ["Heading 1", "Heading 2", "Heading 3"]
for p in doc.paragraphs:
    if p.style.name in styles_to_fix:
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.color.rgb = RGBColor(0, 0, 0)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = (
    "/mnt/c/Users/User/Documents/GitHub/claude-squad/CRUDE_RESILIENCE_v2_OPIM626.docx"
)
doc.save(out_path)
print(f"Saved: {out_path}")
