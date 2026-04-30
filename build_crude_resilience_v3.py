from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)


# ── Helpers ────────────────────────────────────────────────────────────────────
def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pBdr.append(bottom)
    pPr.append(pBdr)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p


def body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
    p.add_run(text)
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
    p.add_run(text)
    return p


def set_font(doc):
    from docx.oxml.ns import qn
    from docx.shared import Pt

    for p in doc.paragraphs:
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(11)


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(6)
r = title.add_run("FRAGILE FLOWS:\nTHE 2026 HORMUZ OIL SHOCK")
r.bold = True
r.font.size = Pt(14)
r.font.name = "Times New Roman"

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(0)
sub.paragraph_format.space_after = Pt(2)
sr = sub.add_run(
    "Supply Chain Vulnerability, Risk, and Resilience\nin the 2026 Iran-Hormuz Crisis"
)
sr.font.size = Pt(12)
sr.font.name = "Times New Roman"

add_hr(doc)

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors.paragraph_format.space_before = Pt(4)
authors.paragraph_format.space_after = Pt(1)
ar = authors.add_run(
    "Group 4: Abhishek Gupta | Almanie Khalid Waleed S | Kushagra Gupta\n"
    "Kaylea TAN Kai Qi | Nippun Aggarwal | Shivam Kumar\n"
    "Ujjwal Sancheti | Tejaswini Vivek"
)
ar.font.size = Pt(10)
ar.font.name = "Times New Roman"

wc = doc.add_paragraph()
wc.alignment = WD_ALIGN_PARAGRAPH.CENTER
wc.paragraph_format.space_before = Pt(2)
wc.paragraph_format.space_after = Pt(2)
wr = wc.add_run(
    "Word count: 4,037  |  OPIM 626 — Risk Management in Global Supply Chains  |  SMU MBA 2026  |  April 2026"
)
wr.font.size = Pt(10)
wr.italic = True
wr.font.name = "Times New Roman"

add_hr(doc)

note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(4)
note.paragraph_format.space_after = Pt(4)
rn = note.add_run(
    "Note: This crisis is ongoing as of April 26, 2026. All firm-level performance data reflects "
    "preliminary estimates. This paper frames itself as a practitioner playbook — a structured set of "
    "diagnostic and response frameworks derived from a live event — not a retrospective case study. "
    "The goal is to extract transferable lessons before the next disruption, not to declare "
    "premature conclusions."
)
rn.italic = True
rn.font.size = Pt(10)
rn.font.name = "Times New Roman"

add_hr(doc)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "EXECUTIVE SUMMARY")

body(
    doc,
    "On 28 February 2026, joint US-Israeli strikes on Iranian military infrastructure triggered "
    "one of the most consequential supply chain disruptions in modern history. Iran's retaliatory "
    "closure of the Strait of Hormuz sent Brent crude surging from approximately $84/bbl to an "
    "intraday peak of $119.50/bbl by 9 March — a 42% spike in under two weeks (BNN Bloomberg, "
    "March 2026). The Baltic Dirty Tanker Index reached a historical high of 3,723 points by "
    "30 March, a 230% year-on-year surge. By 25 March, approximately 1,900 commercial vessels "
    "including 211 crude oil tankers carrying an estimated 190 million barrels remained stranded "
    "in or around the strait (Anadolu Agency, March 2026). The IEA responded with the largest "
    "emergency reserve release in its history — a unanimous 400-million-barrel drawdown across "
    "32 member countries (IEA Executive Director Fatih Birol, CNBC, March 2026).",
)

body(
    doc,
    "Most commentary on the 2026 crisis stops at the price spike. This paper argues that the "
    "crisis is not a price story — it is a supply chain architecture story. The disruption did "
    "not originate where most observers expected, did not operate through the mechanism most "
    "anticipated, and was not overcome by the firms with the most sophisticated trading desks. "
    "It was overcome by firms whose pre-crisis commercial infrastructure was designed with "
    "disruption scenarios in mind.",
)

body(
    doc,
    "This paper introduces the controlled transaction as the central resilience instrument — "
    "a pre-negotiated, pre-authorized commercial arrangement with defined activation triggers "
    "that eliminates the negotiation phase from the crisis response. It traces a four-layer "
    "disruption cascade: war risk insurance withdrawal, tanker routing constraints, refinery "
    "crude slate inflexibility, and information asymmetry. It analyses how capital flow "
    "architecture — the pre-existing structure of correspondent banking, settlement currency "
    "corridors, and cargo tracking interoperability — determined which firms were structurally "
    "positioned to activate controlled transactions and which were not. And it delivers a "
    "practitioner playbook: diagnostic questions, preparation checklists, and response protocols "
    "anchored in the controlled transaction framework.",
)

# ══════════════════════════════════════════════════════════════════════════════
# 1. BACKGROUND AND MOTIVATION
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "1. BACKGROUND AND MOTIVATION")

h2(doc, "1.1  The Crisis")

body(
    doc,
    "The Strait of Hormuz, at its narrowest point just 33 kilometres wide between Oman and "
    "Iran, carries approximately 20-21 million barrels per day of crude oil — roughly 20% "
    "of global seaborne crude trade (EIA, Hormuz Transit Factsheet). Its symbolic significance "
    "as a global energy chokepoint is well-established. What the 2026 crisis revealed is that "
    "the disruption did not originate at the Strait itself — and the actual cascade mechanisms "
    "were financial and logistical, not purely physical.",
)

body(
    doc,
    "The physical closure of the Strait sent shockwaves through global commodity markets. "
    "Brent crude spiked to an intraday peak of $119.50/bbl by 9 March 2026 (BNN Bloomberg, "
    "March 2026). The Baltic Dirty Tanker Index surged to 3,723 points by 30 March, a 230% "
    "year-on-year increase, reflecting the massive rerouting of global tanker capacity "
    "(Baltic Exchange, March 2026). GPS jamming attributed to Iranian electronic warfare "
    "disrupted navigation for over 1,100 vessels in the Persian Gulf (The Signal Group; "
    "Windward, March 2026). J.P. Morgan estimated potential supply losses of up to "
    "4.7 million barrels per day from Iraq and Kuwait alone if southern oilfields were cut "
    "off (Reuters, March 2026). By early April, prices partially stabilised in the "
    "$103-$111/bbl range — not because the crisis had resolved, but because the system "
    "adapted imperfectly, unevenly, and with residual fragility.",
)

h2(doc, "1.2  The Sanctions Foundation")

body(
    doc,
    "The character of the 2026 disruption was shaped by five years of evolving US sanctions "
    "architecture. Following the US withdrawal from the JCPOA in November 2018, Iranian crude "
    "exports collapsed from approximately 2.5 million barrels per day to under 1 mb/d by 2019 "
    "(IEA Oil Market Reports). This collapse restructured Asian refiner exposure asymmetrically.",
)

body(
    doc,
    "Indian state-owned refiners — primarily Indian Oil Corporation (IOC) — systematically "
    "reduced Iranian crude intake from approximately 10% of total imports in 2018 to under "
    "2% by 2025, replaced by Saudi Arab Light, Abu Dhabi crudes, and US crude imports "
    "under long-term contracts. This was primarily a sanctions compliance response — but its "
    "effect was to pre-position Indian refiners closer to alternative supply networks before "
    "the 2026 disruption. Chinese state-owned refiners — Sinopec and CNPC — maintained "
    "higher Iranian exposure through intermediary arrangements, estimated at 5-7% of total "
    "Chinese crude imports in 2025 (Bloomberg data). Japanese refiners reduced Iranian "
    "imports after 2012 and maintained a broadly diversified supplier base anchored on "
    "Saudi Arabia, UAE, Kuwait, and growing US crude imports (US EIA, 2026).",
)

h2(doc, "1.3  Capital Flow Architecture: The Hidden Speed Differentiator")

body(
    doc,
    "A dimension that standard supply chain risk management analysis overlooks is the interaction "
    "between capital flow architecture and physical supply chain pre-positioning. "
    "Geopolitical disruption events operate on two simultaneous timescales: the immediate "
    "physical disruption and a pre-existing capital and trading network architecture that "
    "determines which firms are structurally positioned to respond. Firms embedded in "
    "US-dollar clearing systems and Western commodity trading networks had faster access "
    "to alternative supply chains during the 2026 crisis. Firms operating primarily "
    "through Belt and Road corridors and non-dollar settlement networks faced longer "
    "activation times due to correspondent banking constraints, slower clearance "
    "infrastructure, and less interoperable cargo tracking systems.",
)

body(
    doc,
    "This is a supply chain infrastructure observation, not a political one: correspondent "
    "banking relationships, settlement currency choice, cargo tracking interoperability, "
    "and insurance market membership all reflect pre-existing capital flow architecture "
    "that shapes crisis response speed. The four-layer disruption cascade identified in this "
    "paper is not random in its effects — its impact is filtered through the "
    "pre-existing capital and trading network architecture of each firm.",
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. WHERE VULNERABILITY ACTUALLY MATERIALIZED
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "2. WHERE VULNERABILITY ACTUALLY MATERIALIZED")

body(
    doc,
    "The disruption cascade operated across four distinct layers. Each represents a different "
    "type of supply chain vulnerability, each triggered at a different speed, and each "
    "affected different firms differently depending on their pre-crisis architecture. "
    "Critically, the same four-layer cascade hit firms differently — and capital flow "
    "architecture (Section 1.3) is the primary explanation for why: firms embedded "
    "in dollar-clearing Western networks activated alternative supply faster than those "
    "operating through non-dollar corridors, even when their underlying supplier "
    "relationships were comparable.",
)

h2(doc, "2.1  Layer One: War Risk Insurance Market Freeze")

body(
    doc,
    "The first constraint was financial, not physical. War risk insurance for tankers "
    "operating in or near the Gulf region seized within 48 hours of the escalation. "
    "War risk insurance is underwritten primarily through Lloyd's of London syndicates "
    "and the International Tanker Fuel Insurance Forum (ITFF). When geopolitical risk "
    "elevates, these underwriters reprice rapidly and capacity can withdraw from "
    "specific routing corridors entirely.",
)

body(
    doc,
    "War risk insurance premiums for Gulf routing increased approximately threefold within "
    "the first week, from roughly $0.10-0.15 per barrel to $0.30-0.45 per barrel "
    "(Lloyd's Market Association Commentary, April 3, 2026; Marsh War Risk Insurance "
    "Briefing Note, April 4, 2026). For a VLCC carrying 500,000 barrels, this added "
    "approximately $150,000-$225,000 per voyage — against a cargo value of approximately "
    "$50 million at $100/barrel, a 0.3-0.45% cost increase in isolation. The more "
    "consequential effect was capacity withdrawal: for some vessel operators, no war risk "
    "insurance capacity was available at any price for Gulf routing within the first "
    "7-10 days. This created a functional supply disruption: crude oil was physically "
    "available, but the financial infrastructure to transport it had withdrawn. "
    "This is a tier-3 supply chain vulnerability: an insurance market disruption that "
    "cascaded upstream into physical supply before most firms detected it.",
)

h2(doc, "2.2  Layer Two: Tanker Routing and the Cape Detour")

body(
    doc,
    "The second layer was physical. Vessels avoiding Gulf routing — either because "
    "insurance was unavailable or because operators chose to avoid the premium — rerouted "
    "via the Cape of Good Hope, adding 10-15 days of transit time to Asia-bound routes. "
    "VLCC freight rates for Cape-routed voyages increased by 35-50% compared to "
    "Hormuz-routed equivalents within the first two weeks (Bloomberg Shipping Intelligence, "
    "April 2-8, 2026; Baltic Exchange VLCC assessments, March-April 2026).",
)

body(
    doc,
    "The 10-day transit extension was operationally damaging because most Asian refineries "
    "operate with 5-10 days of crude inventory on hand. A 10-day delay exceeds the "
    "forward planning window of even well-managed refineries. Asian refinery crude inventory "
    "norms range from approximately 5 days for commercial just-in-time operators to "
    "14+ days for refiners with strategic reserve access (IEA data). This means the "
    "rerouting delay was a physical supply threat — not a price problem — and it "
    "affected all Asian refiners regardless of their crude sourcing sophistication. "
    "Firms with inventory buffers absorbed the delay; those without faced throughput cuts. "
    "The capital flow architecture dimension compounded this: firms whose alternative supply "
    "routes required correspondent banking clearance in non-Western currencies faced "
    "additional hours to days of processing time when activating emergency procurement.",
)

h2(doc, "2.3  Layer Three: Refinery Crude Slate Inflexibility")

body(
    doc,
    "The third layer was technical and widely underestimated. A refinery's Nelson "
    "Complexity Index (NCI) determines what crude grades it can process efficiently. "
    "Simple refineries (NCI 4-6) are designed for specific crude types and require "
    "weeks of operational adjustment to process meaningfully different grades. Complex "
    "refineries (NCI 9-12), such as Reliance Industries' Jamnagar complex (NCI 13+), "
    "are designed to accept a wide slate including heavy, light, sweet, and sour crudes.",
)

body(
    doc,
    "Iranian light crude (33-36 degree API, low sulfur) is relatively straightforward "
    "for complex refineries but creates processing challenges for simpler configurations. "
    "Saudi Arab Light (36 degree API, low sulfur) is more interchangeable with Iranian "
    "crude than Russian Urals (32-36 degree API, higher sulfur), which requires "
    "specific desulfurisation capacity. Some firms were switching to grades they could "
    "process efficiently; others were switching to grades that imposed 2-4 weeks of "
    "throughput loss even after securing the cargo. The firms that recovered fastest "
    "were not simply those with alternative supplier relationships — they were those "
    "whose refinery configurations could absorb the grade shift without major "
    "operational disruption.",
)

h2(doc, "2.4  Layer Four: The Visibility Gap")

body(
    doc,
    "The fourth layer was informational. The most sophisticated refiners — led by "
    "Reliance Industries' trading desk — monitored real-time vessel movements via AIS "
    "(Automatic Identification System) tracking platforms (Kpler, MarineTraffic). These "
    "systems allow traders to see vessels diverting from Hormuz to Cape routing in "
    "near-real-time: typically 24-48 hours before official port or shipping reports "
    "confirm the shift.",
)

body(
    doc,
    "Indian Oil Corporation, as a state-owned enterprise with strategic inventory mandates, "
    "had access to government intelligence and logistics coordination channels that provided "
    "earlier warning than commercial data alone. Chinese state-owned refiners had access "
    "to similar state channels but faced decision latency from multi-layer NOC governance "
    "structures — the same approval chains that constrain commercial flexibility in normal "
    "times slowed crisis response even when early intelligence was available. "
    "Choi, Rogers, and Vakil (2020) argue that the fundamental supply chain "
    "vulnerability in modern commerce is opacity into tier-2 and tier-3 supply chains. "
    "The 2026 Iran crisis validates this: the visibility gap — which firms had real-time "
    "tier-2 (shipping) and tier-3 (insurance) intelligence and which did not — "
    "translated directly into a 24-72 hour decision advantage for firms with "
    "monitoring systems in place.",
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. THE CONTROLLED TRANSACTION: THE RESILIENCE INSTRUMENT THAT MATTERED
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "3. THE CONTROLLED TRANSACTION: THE RESILIENCE INSTRUMENT THAT MATTERED")

body(
    doc,
    "The empirical observation from the crisis is straightforward: firms with pre-established "
    "alternative supplier relationships recovered faster than firms that had to negotiate "
    "new arrangements under pressure. But this observation conceals an important "
    "distinction. Having an alternative supplier relationship is not the same as having "
    "a controlled transaction.",
)

h2(doc, "3.1  Definition")

body(
    doc,
    "A controlled transaction is a pre-negotiated, pre-authorized commercial arrangement "
    "between counterparties that specifies: (1) defined activation triggers — the crisis "
    "conditions under which the arrangement will be activated; (2) agreed volumes — "
    "minimum and maximum quantities deliverable under activation; (3) pre-agreed pricing "
    "mechanics — a formula or price band that applies upon activation, negotiated before "
    "the crisis; (4) pre-authorized decision authority — specified individuals or roles "
    "with the commercial authority to activate the transaction without additional "
    "approval chains.",
)

body(
    doc,
    "The critical feature of a controlled transaction is that it eliminates the negotiation "
    "phase from the crisis response. Normal commercial relationships require re-negotiation "
    "under crisis conditions: volumes must be confirmed, prices must be agreed, terms must "
    "be documented, and approvals must be obtained — all while the disruption is already "
    "cascading. A controlled transaction pre-resolves these steps. The counterparties "
    "have agreed in advance: if X happens, we execute Y on these terms, and Z has the "
    "authority to trigger it.",
)

h2(doc, "3.2  India and China: Two Models of Supply Security")

body(
    doc,
    "Indian Oil Corporation and Reliance Industries had pre-established term contracts "
    "with Saudi Aramco, Abu Dhabi National Oil Company (ADNOC), and US crude exporters. "
    "These were not emergency spot purchases negotiated in March 2026 — they were existing "
    "commercial arrangements that included minimum volume commitments, pre-agreed pricing "
    "formulas, and commercial decision authority that allowed trading teams to execute "
    "within hours of crisis activation. This is the controlled transaction in practice.",
)

body(
    doc,
    "Chinese NOCs — Sinopec and CNPC — had long-term Iranian supply agreements "
    "structured with take-or-pay provisions: payment obligations continued regardless of "
    "whether Chinese buyers took delivery. These arrangements optimised for supply security "
    "and relationship maintenance with Iran — not for activation flexibility. When the "
    "2026 disruption hit, Chinese NOCs faced a structurally different problem: locked "
    "into Iranian supply commitments that could not be quickly redirected, while "
    "simultaneously needing to secure alternative supply through commercial channels "
    "that required new negotiations. Switching to Russian crude demanded extended "
    "commercial discussions and government approvals for volume reallocations — a "
    "process measured in weeks, not hours.",
)

body(
    doc,
    "The controlled transaction lens clarifies the India-China divergence in a way that "
    "the generic 'supplier diversification' framing does not. The question is not simply "
    "whether a firm had alternative suppliers — it is whether those alternatives were "
    "structured as controlled transactions with pre-authorised activation. IOC's "
    "alternative supply relationships were. CNPC's were not. And the capital flow "
    "architecture context deepened this divide: activating alternative supply through "
    "dollar-denominated Western contracts required fewer correspondent banking steps "
    "for IOC than activating alternative supply through non-dollar corridors "
    "for CNPC — an infrastructure friction that added hours to days to "
    "the Chinese NOC response.",
)

h2(doc, "3.3  The Limits of Hedging: Price Instruments vs. Availability Instruments")

body(
    doc,
    "Sodhi and Tang (2012) categorise supply chain risks into disruption risks, "
    "supply risks, and demand risks. Each requires different mitigation instruments. "
    "The 2026 Iran-Hormuz crisis was unambiguously a disruption risk — sudden, "
    "external, availability-threatening — not a price risk.",
)

body(
    doc,
    "Chinese NOCs are sophisticated hedgers. Sinopec and CNPC use Brent futures, "
    "options, and swaps extensively to manage price risk. These instruments performed "
    "exactly as designed during the 2026 crisis: they protected against the cost of "
    "Brent spiking to $100+. But they did not protect against physical supply "
    "disruption — the scenario where alternative crude was available but at a "
    "$15-20/barrel premium with 10 days of additional transit time. "
    "Hedging instruments address price risk. Controlled transactions, redundancy "
    "(inventory buffers), and flexibility (refinery configuration) are the appropriate "
    "instruments for disruption risks. Chinese NOCs were well-hedged against the "
    "wrong risk type — which is why their hedges provided financial protection "
    "without preventing physical supply disruption. This is the critical "
    "distinction Sodhi and Tang's framework illuminates.",
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. FRAMEWORK APPLICATION
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "4. FRAMEWORK APPLICATION")

h2(doc, "4.1  Sheffi: Five Dimensions Mapped to the Crisis")

body(
    doc,
    "Sheffi (2005) identifies five dimensions of organisational resilience: redundancy, "
    "flexibility, awareness, adaptability, and organisational capabilities. "
    "The 2026 Iran-Hormuz crisis provides an empirical test of each:",
)

tbl = doc.add_table(rows=6, cols=3)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
hdr[0].text = "Sheffi Dimension"
hdr[1].text = "Crisis Observation"
hdr[2].text = "Firms That Benefited"
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    (
        "Redundancy",
        "14+ day inventory buffers absorbed the 10-day Cape rerouting delay without emergency spot procurement",
        "IOC (strategic reserves); Japanese refiners (national security inventories)",
    ),
    (
        "Flexibility",
        "NCI 13+ refinery configurations processed grade-switched crude without major throughput loss; simpler refineries faced 2-4 weeks of reduced throughput",
        "Reliance Industries (Jamnagar); Japanese refiners with complex configurations",
    ),
    (
        "Awareness",
        "Real-time AIS vessel tracking provided 24-48 hours of early warning before physical supply impact reached refinery gates",
        "Reliance trading desk; IOC (government intelligence channels)",
    ),
    (
        "Adaptability",
        "Pre-established alternative supplier relationships (Saudi Aramco, ADNOC, US crude) activated without new negotiations in 3-6 days",
        "IOC; Japanese refiners with diversified term contracts",
    ),
    (
        "Organisational Capabilities",
        "Pre-authorised decision authority enabled supplier switches within 24-72 hours without multi-week approval chains",
        "Reliance trading desk; IOC (pre-authorised crisis protocols)",
    ),
]
for i, (dim, obs, firms) in enumerate(rows, 1):
    r = tbl.rows[i].cells
    r[0].text = dim
    r[1].text = obs
    r[2].text = firms

doc.add_paragraph()

h2(doc, "4.2  Sodhi and Tang: Risk Taxonomy Applied")

bullet(
    doc,
    "Disruption risks (sudden, external, potentially extended-duration) require redundancy "
    "and flexibility. The Iran-Hormuz crisis is unambiguously a disruption risk in this "
    "framework. Mitigation instruments: inventory buffers, flexible refinery configurations, "
    "controlled transactions.",
    bold_prefix="Disruption risk:",
)
bullet(
    doc,
    "Supply risks (uncertainty in quality, quantity, or timing of supply from a known "
    "supplier) require supplier management and contractual mechanisms. Relevant to the "
    "refinery crude slate problem — switching to the wrong crude grade creates a supply "
    "quality risk that manifests as reduced throughput. Not the primary risk type here.",
    bold_prefix="Supply risk:",
)
bullet(
    doc,
    "Demand risks are not directly relevant to this crisis, though downstream product "
    "demand destruction could emerge if refinery throughput cuts propagate into "
    "refined product shortages.",
    bold_prefix="Demand risk:",
)

body(
    doc,
    "The taxonomy explains why Chinese NOC hedging strategies — which addressed price "
    "risk — did not protect against the disruption risk they faced. The risk type "
    "and the mitigation instrument must match. Firms that used only price instruments "
    "suffered physical disruption even as their financial hedges performed.",
)

h2(doc, "4.3  Choi, Rogers, and Vakil: Visibility Applied")

body(
    doc,
    "Choi, Rogers, and Vakil (2020) argue that the post-COVID supply chain vulnerability "
    "is opacity into tier-2 and tier-3 supply chains. In the oil refining supply chain: "
    "tier-1 is the crude supplier; tier-2 is the shipping and logistics network; "
    "tier-3 is the financial infrastructure (insurance, banking, payment clearing). "
    "The geopolitical environment — sanctions regimes, diplomatic relationships, and "
    "cross-border payment restrictions — operates above these tiers as the structural "
    "context that shapes tier-3 and tier-2 behavior simultaneously.",
)

body(
    doc,
    "The visibility gap in the 2026 crisis operated at tier-2 and tier-3. Firms did not "
    "primarily lack visibility into their crude suppliers — they knew who their suppliers "
    "were. They lacked visibility into whether vessels carrying their crude could be "
    "insured, whether they would reroute, and when the rerouting would translate into "
    "physical supply delays. Firms with real-time AIS tracking detected tier-2 disruption "
    "24-48 hours before those without. Firms monitoring war risk insurance premium "
    "movements detected tier-3 disruption before it manifested in physical routing changes.",
)
body(
    doc,
    "This 24-72 hour early warning window mattered because it determined whether a firm "
    "could activate its controlled transaction while inventory was still intact — or only "
    "after throughput cuts had already begun. A refinery running 10 days of crude inventory "
    "has no spare time: if the activation decision arrives on day 10, the controlled "
    "transaction activates on day 10 and crude arrives on day 20, creating a 10-day gap. "
    "If the same decision arrives on day 8, the gap closes to 2 days. The visibility "
    "instrument does not itself resolve the supply problem — the controlled transaction "
    "does. But without the early warning, the controlled transaction arrives too late "
    "to prevent disruption. The firms without visibility were already behind when they "
    "started responding.",
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. THE PLAYBOOK
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "5. THE PLAYBOOK: A STRUCTURED GUIDE FOR THE NEXT CRISIS")

body(
    doc,
    "The playbook moves from diagnostic questions through preparation checklists to "
    "response protocols, closing with the structural uncertainties that could reprice "
    "any specific set of recommendations.",
)

h2(doc, "5.1  Diagnostic Questions: Where Does Your Vulnerability Live?")

bullet(
    doc,
    "Insurance layer: What is your current war risk insurance coverage for Gulf routing? "
    "What premium level makes continued Gulf routing economically prohibitive for your "
    "vessel operators — and what is your contingency routing plan?",
    bold_prefix="Insurance layer:",
)
bullet(
    doc,
    "Routing layer: What is your current crude inventory in days of forward supply? "
    "If your primary supply route added 10 days of transit time tomorrow, would you "
    "face runout risk?",
    bold_prefix="Routing layer:",
)
bullet(
    doc,
    "Refinery input layer: Which crude grades can your refineries process at near-full "
    "throughput? How many days of operational adjustment would a grade switch require?",
    bold_prefix="Refinery input layer:",
)
bullet(
    doc,
    "Capital flow layer: Through which correspondent banking and settlement currency "
    "corridors does your emergency procurement activate? How many hours to days of "
    "processing friction does your banking infrastructure add to emergency activation?",
    bold_prefix="Capital flow layer:",
)
bullet(
    doc,
    "Visibility layer: Do you have real-time monitoring of AIS vessel movements for "
    "your in-transit cargoes? Who in your organisation sees this data, and what is "
    "their authorisation when they see it?",
    bold_prefix="Visibility layer:",
)

h2(doc, "5.2  Preparation Checklist: Building Controlled Transaction Infrastructure")

body(
    doc,
    "Building controlled transaction infrastructure requires four parallel workstreams:",
)

bullet(
    doc,
    "Supplier layer: For each major crude grade you process, identify three alternative "
    "suppliers. Negotiate baseline contracts at minimum viable volumes. Include defined "
    "activation triggers, pre-agreed pricing formulas, and pre-authorised decision "
    "authority in the contract structure. Without these three elements, the contract is "
    "not a controlled transaction — it is a theoretical alternative that cannot be "
    "exercised under crisis pressure.",
    bold_prefix="Supplier layer:",
)
bullet(
    doc,
    "Inventory layer: IEA data supports 14+ days of crude inventory as adequate buffer "
    "for a 10-14 day disruption event. Below 10 days is a material vulnerability. "
    "Inventory carries a working capital cost — but it is the most direct redundancy "
    "instrument for a disruption risk.",
    bold_prefix="Inventory layer:",
)
bullet(
    doc,
    "Refinery layer: Audit your refinery's actual crude grade flexibility using Nelson "
    "Complexity Index and crude assay data. If your refinery is configured for specific "
    "crude types, your controlled transaction strategy must be limited to alternative "
    "suppliers delivering the grades your refinery can process.",
    bold_prefix="Refinery layer:",
)
bullet(
    doc,
    "Capital flow layer: Map the correspondent banking steps required to activate emergency "
    "procurement in each alternative currency corridor. Pre-position correspondent "
    "banking relationships for your contingency suppliers. Every hour of banking "
    "friction in your activation path is latency you cannot afford in a crisis.",
    bold_prefix="Capital flow layer:",
)
bullet(
    doc,
    "Organisational layer: Establish and document pre-authorised decision thresholds "
    "for crisis supplier activation. Specify who can activate a controlled transaction, "
    "at what volume threshold, and under what trigger conditions. Practice these "
    "protocols in annual crisis simulations.",
    bold_prefix="Organisational layer:",
)

h2(doc, "5.3  Response Protocol: The First 72 Hours")

bullet(
    doc,
    "Hour 0-24: Assess. Confirm the nature of the disruption: insurance withdrawal, "
    "routing change, or physical blockage? Activate the crisis team. Pull real-time "
    "AIS tracking for all in-transit cargoes. Quantify your inventory runway in days. "
    "Identify which of your five vulnerability layers is the binding constraint.",
    bold_prefix="Hour 0-24:",
)
bullet(
    doc,
    "Hour 24-72: Activate. Activate pre-established controlled transactions for the "
    "affected supply layer. Draw down inventory to extend supply runway while "
    "alternative arrangements are activated. Brief treasury on currency and hedging "
    "implications of alternative supplier procurement.",
    bold_prefix="Hour 24-72:",
)
bullet(
    doc,
    "Day 3-14: Manage. Assess likely duration: temporary (under 30 days) or "
    "structural? For temporary disruptions: maintain controlled transaction activation, "
    "do not restructure the supplier base permanently. For structural disruptions: "
    "execute formal supplier base restructuring and evaluate take-or-pay renegotiation.",
    bold_prefix="Day 3-14:",
)

body(
    doc,
    "Oil-specific early warning signals to monitor continuously: VLCC routing patterns "
    "(AIS rerouting signals disruption onset); war risk insurance premium levels "
    "(spikes above 50% week-over-week signal market stress); Brent backwardation "
    "(front-month premium over 6-month of more than $3-4/barrel signals near-term "
    "supply tightness); VLCC spot rate differentials between Cape and Hormuz routing "
    "(differential widening signals rerouting cost repricing).",
)

h2(doc, "5.4  Structural Uncertainty: The Saudi-Iran Normalisation Scenario")

body(
    doc,
    "The playbook above is built for the disruption scenario that materialised in "
    "March 2026. A competent risk management analysis must address the scenario that "
    "could reprice its recommendations: Saudi-Iran diplomatic normalisation.",
)

body(
    doc,
    "Chinese-mediated Saudi-Iran engagement has been a live geopolitical dynamic since "
    "2021, culminating in the 2023 Beijing Accord. A future normalisation — potentially "
    "involving US-Iran nuclear negotiations or Chinese-brokered regional détente — "
    "could lift secondary sanctions and return Iranian crude to global markets at scale "
    "(potentially 2+ mb/d of additional supply). This would reprice Middle Eastern "
    "crude benchmarks and disrupt the alternative supply relationships activated "
    "during the 2026 crisis.",
)

body(
    doc,
    "The resilience infrastructure built for the disruption scenario must be designed "
    "with scenario-specific activation terms and periodic review triggers. Specifically: "
    "the alternative supplier contracts activated in 2026 should include sunset provisions "
    "tied to Iranian crude returning to market. The practitioner message must be nuanced: "
    "resilience architecture should be robust across multiple scenarios, not optimised "
    "exclusively for the most recent disruption. The firms that will manage the next "
    "disruption best are those that build scenario flexibility into their supply chain "
    "infrastructure — not those that optimise for the last crisis. This is not a price "
    "story. It is a structure story.",
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "6. CONCLUSION")

body(
    doc,
    "The 2026 Iran-Hormuz crisis is ongoing. Final recovery outcomes for Indian, "
    "Chinese, and Japanese refiners are not yet established. This paper has not declared "
    "conclusions — it has offered a diagnostic and a playbook.",
)

body(
    doc,
    "The diagnostic is this: the disruption did not operate through the mechanism "
    "most observers expected. The Strait of Hormuz was not physically blocked. "
    "The disruption originated in the financial infrastructure of oil transport — the "
    "war risk insurance market — and propagated through tanker routing logistics, "
    "refinery technical constraints, information asymmetry, and capital flow "
    "architecture. The four-layer cascade was filtered through the pre-existing "
    "capital and trading network structure of each firm.",
)

body(
    doc,
    "The central instrument is the controlled transaction: pre-negotiated, "
    "pre-authorised commercial arrangements with defined activation triggers and "
    "decision authority. Sheffi's five dimensions tell you where to invest. "
    "Sodhi and Tang's taxonomy tells you which risk type each investment "
    "addresses. Choi et al.'s visibility framework tells you where to look "
    "when the next crisis begins. The capital flow architecture lens tells you "
    "that the speed of your activation is not only organisational — it is "
    "infrastructural. The correspondent banking corridors and settlement "
    "currency architecture through which your emergency procurement travels "
    "are as much a part of your resilience infrastructure as your "
    "refinery configuration or inventory buffer.",
)

body(
    doc,
    "The structural uncertainty is this: Saudi-Iran normalisation is a credible "
    "scenario that would reprice the investments made during the 2026 crisis. "
    "Resilience architecture designed for a single scenario is fragile architecture. "
    "The firms that will manage the next disruption best build scenario flexibility "
    "into their supply chain infrastructure — not those that optimise for "
    "the most recent crisis. This is a supply chain risk management story, "
    "not a commodity trading story.",
)

# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, "REFERENCES")

refs = [
    (
        "Course Readings:",
        [
            "Choi, T.Y., Rogers, D.S., and Vakil, B. (2020). 'Coronavirus is a Wake-Up "
            "Call for Supply Chain Management.' Harvard Business Review, April 2020.",
            "Sheffi, Y. (2005). The Resilient Enterprise: Overcoming Vulnerability for "
            "Competitive Advantage. MIT Press.",
            "Sodhi, M.S., and Tang, C.S. (2012). Managing Supply Chain Risks. Springer.",
        ],
    ),
    (
        "Data Sources:",
        [
            "Anadolu Agency. (2026). 'Strait of Hormuz: 1,900 Vessels Stranded as Tensions "
            "Escalate,' March 2026.",
            "Baltic Exchange. (2026). VLCC freight rate assessments, BDTI data, March-April 2026. "
            "https://www.balticexchange.com/",
            "Bloomberg L.P. (2026). Brent crude daily pricing, VLCC freight indices, war risk "
            "insurance premium data, March-April 2026. Bloomberg Terminal subscription.",
            "BNN Bloomberg. (2026). Oil market reporting, March 2026.",
            "CNBC. (2026). 'IEA Launches Largest Emergency Reserve Release in History,' March 2026. "
            "Statement by IEA Executive Director Fatih Birol.",
            "IEA. (2026). Monthly Oil Market Reports, March-April 2026. https://www.iea.org/",
            "Kpler and MarineTraffic. (2026). Tanker tracking and AIS vessel movement data, "
            "March-April 2026. Subscription data accessed April 2026.",
            "Lloyd's Market Association. (2026). War Risk Insurance Market Commentary, "
            "April 3, 2026. https://www.lloyds.com/",
            "Marsh and McLennan. (2026). War Risk Insurance Briefing Note, April 4, 2026.",
            "OPEC. (2026). Monthly Oil Market Reports, February-March 2026. https://www.opec.org/",
            "Reuters. (2026). 'J.P. Morgan Estimates Potential Supply Losses of 4.7 mb/d,' "
            "March 2026.",
            "The Signal Group and Windward. (2026). GPS jamming and vessel navigation "
            "disruption reporting, March 2026.",
            "US EIA. (2026). Oil Market Reports, Hormuz Transit Factsheet, March-April 2026. "
            "https://www.eia.gov/",
        ],
    ),
]

for category, items in refs:
    ph = doc.add_paragraph()
    ph.paragraph_format.space_before = Pt(6)
    ph.paragraph_format.space_after = Pt(2)
    rh = ph.add_run(category)
    rh.bold = True
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

# ── Font application ─────────────────────────────────────────────────────────────
set_font(doc)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = (
    "/mnt/c/Users/User/Documents/GitHub/claude-squad/FRAGILE_FLOWS_Writeup_Group4.docx"
)
doc.save(out_path)
print(f"Saved: {out_path}")
