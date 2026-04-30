from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1.25)


# ── Helpers ────────────────────────────────────────────────────────────────────
def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pBdr.append(bottom)
    pPr.append(pBdr)


def body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
    p.add_run(text)
    return p


def run_font(run, size=11, name="Times New Roman"):
    run.font.size = Pt(size)
    run.font.name = name


# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(4)
r = title.add_run("FRAGILE FLOWS:\nTHE 2026 HORMUZ OIL SHOCK")
r.bold = True
r.font.size = Pt(14)
r.font.name = "Times New Roman"
run_font(r, 14)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(0)
sub.paragraph_format.space_after = Pt(2)
sr = sub.add_run(
    "Supply Chain Vulnerability, Risk, and Resilience\nin the 2026 Iran-Hormuz Crisis"
)
sr.font.size = Pt(12)
sr.font.name = "Times New Roman"

add_hr(doc)

# ── Authors block ─────────────────────────────────────────────────────────────
authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors.paragraph_format.space_before = Pt(4)
authors.paragraph_format.space_after = Pt(1)
ar = authors.add_run(
    "Group 4\n"
    "Abhishek Gupta  |  Almanie Khalid Waleed S  |  Kushagra Gupta\n"
    "Kaylea TAN Kai Qi  |  Nippun Aggarwal  |  Shivam Kumar\n"
    "Ujjwal Sancheti  |  Tejaswini Vivek"
)
ar.font.size = Pt(10)
ar.font.name = "Times New Roman"

wc = doc.add_paragraph()
wc.alignment = WD_ALIGN_PARAGRAPH.CENTER
wc.paragraph_format.space_before = Pt(2)
wc.paragraph_format.space_after = Pt(2)
wr = wc.add_run("Word count: 1,512  |  Expandable to: 4,000 words")
wr.font.size = Pt(10)
wr.font.name = "Times New Roman"
wr.italic = True

add_hr(doc)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 1. BACKGROUND AND MOTIVATION
# ══════════════════════════════════════════════════════════════════════════════
h1_p = doc.add_paragraph("1.  Background and Motivation")
h1_p.runs[0].bold = True
h1_p.runs[0].font.size = Pt(12)
h1_p.runs[0].font.name = "Times New Roman"
h1_p.paragraph_format.space_before = Pt(8)
h1_p.paragraph_format.space_after = Pt(4)

body(
    doc,
    "On 28 February 2026, joint US-Israeli strikes on Iranian military infrastructure "
    "triggered one of the most consequential supply chain disruptions in modern history. "
    "Iran's retaliatory closure of the Strait of Hormuz sent Brent crude surging from "
    "approximately $84/bbl to an intraday peak of $119.50/bbl by 9 March — a 42% spike "
    "in under two weeks (BNN Bloomberg, March 2026). The Baltic Dirty Tanker Index "
    "reached a historical high of 3,723 points by 30 March, a 230% year-on-year surge, "
    "while GPS jamming attributed to Iranian electronic warfare disrupted navigation for "
    "over 1,100 vessels in the Persian Gulf (The Signal Group; Windward, March 2026). "
    "By 25 March, approximately 1,900 commercial vessels, including 211 crude oil tankers "
    "carrying an estimated 190 million barrels, remained stranded in or around the strait "
    "(Anadolu Agency, March 2026).",
)

body(
    doc,
    "The disruption's scale is without modern precedent. J.P. Morgan estimated potential "
    "supply losses of up to 4.7 million barrels per day from Iraq and Kuwait alone "
    "(Reuters, March 2026). The IEA responded with the largest emergency reserve release "
    "in its history — a unanimous 400-million-barrel drawdown across 32 member countries "
    "(IEA Executive Director Fatih Birol, CNBC, March 2026). By early April, prices "
    "partially stabilised in the $103-$111/bbl range, not because the crisis had "
    "resolved, but because the system adapted imperfectly, unevenly, and with residual fragility.",
)

body(
    doc,
    "Most commentary on the 2026 crisis stops at the price spike. This is insufficient "
    "for supply chain managers who must make sourcing, routing, and inventory decisions "
    "under real uncertainty. This paper argues that the crisis is not a price story — "
    "it is a supply chain architecture story. The disruption did not originate where "
    "most observers expected, did not operate through the mechanism most anticipated, "
    "and was not overcome by the firms with the most sophisticated trading desks. "
    "It was overcome by firms whose pre-crisis commercial infrastructure was designed "
    "with disruption scenarios in mind. This paper extracts the transferable lessons.",
)


# ══════════════════════════════════════════════════════════════════════════════
# 2. SCOPE AND RESEARCH QUESTIONS
# ══════════════════════════════════════════════════════════════════════════════
h2_p = doc.add_paragraph("2.  Scope and Research Questions")
h2_p.runs[0].bold = True
h2_p.runs[0].font.size = Pt(12)
h2_p.runs[0].font.name = "Times New Roman"
h2_p.paragraph_format.space_before = Pt(8)
h2_p.paragraph_format.space_after = Pt(4)

body(
    doc,
    "The original proposal examined four questions across broad stakeholder groups. "
    "This paper narrows scope to one central argument with one overriding research "
    "question — the rest follow from it.",
)

body(
    doc,
    "Central argument: The 2026 crisis is not a price story. It is a story about "
    "which pre-crisis supply chain architecture decisions determined which firms "
    "recovered fast and which did not. The specific instrument that separated fast "
    "recoverers from slow was the controlled transaction — a pre-negotiated, "
    "pre-authorized commercial arrangement with defined activation triggers that "
    "eliminates the negotiation phase from the crisis response. Firms with "
    "controlled transactions recovered in days. Firms without them recovered in weeks.",
)

body(doc, "Supporting research questions:", bold_prefix="")

body(
    doc,
    "Where did the disruption cascade originate and through which layers did it propagate? "
    "Beyond the Strait of Hormuz itself, the paper identifies four distinct vulnerability "
    "layers: war risk insurance, tanker routing, refinery crude slate inflexibility, "
    "and information asymmetry. Each operated at a different speed and affected different "
    "firms differently.",
    bold_prefix="RQ1.",
)
body(
    doc,
    "What is a controlled transaction and why did it separate fast recoverers from slow? "
    "The paper introduces the controlled transaction as the central resilience instrument, "
    "distinguishes it from hedging (a price instrument) and from political relationships "
    "(a strategic access instrument), and traces how it operated differently for Indian "
    "and Chinese refiners.",
    bold_prefix="RQ2.",
)
body(
    doc,
    "What structural changes should an oil-dependent refiner make before the next "
    "geopolitical crisis? The paper delivers a practitioner playbook: a structured "
    "set of diagnostic questions, preparation checklists, and response protocols "
    "anchored in the controlled transaction framework and validated against the "
    "2026 observations.",
    bold_prefix="RQ3.",
)


# ══════════════════════════════════════════════════════════════════════════════
# 3. ANALYTICAL APPROACH
# ══════════════════════════════════════════════════════════════════════════════
h3_p = doc.add_paragraph("3.  Analytical Approach")
h3_p.runs[0].bold = True
h3_p.runs[0].font.size = Pt(12)
h3_p.runs[0].font.name = "Times New Roman"
h3_p.paragraph_format.space_before = Pt(8)
h3_p.paragraph_format.space_after = Pt(4)

body(
    doc,
    "The analysis proceeds in three stages, each corresponding to one research question.",
    bold_prefix="Stage 1 — Vulnerability Cascade Mapping (RQ1):",
)

body(
    doc,
    "The supply network is mapped across three layers: upstream production (Iran, "
    "Saudi Arabia, UAE), seaborne transport (VLCC routes via Hormuz, Suez Canal, "
    "and Cape of Good Hope), and downstream refining hubs in India, China, and Japan. "
    "The analysis identifies the four disruption layers beyond the Strait itself: "
    "war risk insurance concentration, tanker fleet availability, refinery crude slate "
    "inflexibility, and tier-2/tier-3 visibility gaps. The cascade mechanism — how "
    "disruption propagated from insurance withdrawal through routing constraints, "
    "to refinery inputs, to decision latency — is traced chronologically using AIS "
    "tracking data, Lloyd's insurance market reports, and VLCC freight rate data.",
)

body(
    doc,
    "A critical dimension that standard SCRM analysis overlooks is the interaction between "
    "capital flow architecture and physical supply chain pre-positioning. "
    "Geopolitical disruption events operate on two simultaneous timescales: "
    "the immediate physical disruption and a pre-existing capital and trading network "
    "architecture that determines which firms are structurally positioned to respond. "
    "Firms embedded in US-dollar clearing systems and Western commodity trading networks "
    "had faster access to alternative supply chains during the 2026 crisis; "
    "firms operating primarily through Belt and Road corridors and non-dollar "
    "settlement networks faced longer activation times due to correspondent banking "
    "constraints, slower clearance infrastructure, and less interoperable "
    "cargo tracking systems. This is a supply chain infrastructure observation: "
    "correspondent banking relationships, settlement currency choice, cargo tracking "
    "interoperability, and insurance market membership all reflect pre-existing "
    "capital flow architecture that shapes crisis response speed. "
    "The four-layer cascade is not random in its effects; its impact "
    "is filtered through the pre-existing capital and trading network architecture "
    "of each firm.",
    bold_prefix="Capital Flow Architecture:",
)

body(
    doc,
    "The risk type is classified using Sodhi and Tang's taxonomy: the 2026 Iran-Hormuz "
    "crisis is unambiguously a disruption risk (sudden, external, availability-threatening) "
    "rather than a supply risk or demand risk. This classification drives the instrument "
    "analysis in Stage 2.",
    bold_prefix="Stage 2 — Controlled Transaction Analysis (RQ2):",
)

body(
    doc,
    "The paper introduces the controlled transaction as the central analytical construct: "
    "a pre-negotiated, pre-authorized commercial arrangement specifying (1) defined "
    "activation triggers, (2) agreed volumes, (3) pre-agreed pricing mechanics, and "
    "(4) pre-authorized decision authority. This stage applies the construct to the "
    "divergent responses of Indian Oil Corporation and Reliance Industries (which "
    "had controlled transactions with Saudi Aramco, ADNOC, and US crude suppliers) "
    "versus Sinopec and CNPC (which had long-term Iranian supply agreements "
    "structured as take-or-pay arrangements, the inverse of a controlled transaction "
    "for a disruption scenario). The analysis distinguishes controlled transactions "
    "from hedging instruments using Sodhi and Tang's framework: hedges address price "
    "risk; controlled transactions address availability risk.",
)

body(
    doc,
    "Sheffi's five resilience dimensions (redundancy, flexibility, awareness, "
    "adaptability, organizational capabilities) are mapped to the crisis observations "
    "and the controlled transaction framework. Choi, Rogers, and Vakil's visibility "
    "model is applied to explain which firms had 24-48 hours of early warning "
    "and which did not.",
    bold_prefix="Stage 3 — Practitioner Playbook (RQ3):",
)

body(
    doc,
    "The final stage synthesises the analytical findings into a structured practitioner "
    "playbook: (a) diagnostic questions identifying which of the four disruption layers "
    "is the binding constraint for a given firm; (b) a preparation checklist structured "
    "around the controlled transaction framework; (c) a 72-hour crisis response protocol "
    "with oil-specific early warning indicators; and (d) a structural uncertainty "
    "discussion addressing the Saudi-Iran diplomatic normalisation scenario as a "
    "repricing risk for any resilience investments made during the 2026 crisis. "
    "The playbook is validated against the observed responses of Indian, Chinese, "
    "and Japanese refiners.",
)


# ══════════════════════════════════════════════════════════════════════════════
# 4. DATA SOURCES
# ══════════════════════════════════════════════════════════════════════════════
h4_p = doc.add_paragraph("4.  Data Sources")
h4_p.runs[0].bold = True
h4_p.runs[0].font.size = Pt(12)
h4_p.runs[0].font.name = "Times New Roman"
h4_p.paragraph_format.space_before = Pt(8)
h4_p.paragraph_format.space_after = Pt(4)

body(
    doc,
    "Physical flows and shipping data: Kpler, MarineTraffic, Baltic Exchange "
    "(BDTI, VLCC charter rates), March-April 2026.",
    bold_prefix="Shipping and logistics:",
)
body(
    doc,
    "Market and pricing data: Bloomberg L.P. (Brent crude daily pricing, VLCC freight "
    "indices, war risk insurance premium data), Refinitiv Eikon, US EIA Weekly "
    "Petroleum Status Reports, IEA Oil Market Reports, OPEC Monthly Reports, "
    "March-April 2026.",
    bold_prefix="Market and pricing:",
)
body(
    doc,
    "Insurance data: Lloyd's Market Association War Risk Insurance Commentary "
    "(April 3, 2026), Marsh & McLennan War Risk Insurance Briefing Note "
    "(April 4, 2026).",
    bold_prefix="Insurance:",
)
body(
    doc,
    "Company-level responses: Indian Oil Corporation Q1 2026 Earnings Call "
    "(April 15, 2026), Sinopec Group Q1 2026 Earnings Guidance (April 2026), "
    "Reuters reporting on Refinitiv and Bloomberg-sourced crude procurement data "
    "(March 2026).",
    bold_prefix="Corporate disclosures:",
)
body(
    doc,
    "Capital flow and network architecture analysis: supply chain pre-positioning "
    "context derived from the capital flow architecture discussion above, "
    "synthesised from publicly available analysis of correspondent banking patterns, "
    "dollar and non-dollar settlement corridors, and commodity trading network structure.",
    bold_prefix="Capital flow analysis:",
)
body(
    doc,
    "Policy responses: IEA and OPEC statements, March-April 2026; government "
    "communications from India, China, and the United States.",
    bold_prefix="Policy:",
)


# ══════════════════════════════════════════════════════════════════════════════
# 5. EXPECTED CONTRIBUTION
# ══════════════════════════════════════════════════════════════════════════════
h5_p = doc.add_paragraph("5.  Expected Contribution")
h5_p.runs[0].bold = True
h5_p.runs[0].font.size = Pt(12)
h5_p.runs[0].font.name = "Times New Roman"
h5_p.paragraph_format.space_before = Pt(8)
h5_p.paragraph_format.space_after = Pt(4)

body(doc, "This paper makes three contributions to SCRM literature and practice.")

body(
    doc,
    "First, it introduces the controlled transaction as a named resilience instrument "
    "for disruption risks, grounded in Sodhi and Tang's risk taxonomy and operationalised "
    "through the 2026 Iran-Hormuz observations. This fills a gap between the academic "
    "framework (which identifies risk types) and the practitioner toolkit (which needs "
    "specific instruments for each risk type).",
    bold_prefix="Theoretical:",
)
body(
    doc,
    "Second, it traces a four-layer disruption cascade — war risk insurance, tanker "
    "routing, refinery crude slate, and information asymmetry — in a single crisis event, "
    "demonstrating how these layers interact non-linearly and how firm-level architecture "
    "determines which layers become binding constraints.",
    bold_prefix="Empirical:",
)
body(
    doc,
    "Third, it produces a practitioner-relevant playbook — diagnostic questions, preparation "
    "checklists, and response protocols — validated against real operational decisions. "
    "The playbook is designed for procurement executives and risk officers at "
    "oil-dependent firms, applicable to the next geopolitical supply shock "
    "regardless of geography or asset class.",
    bold_prefix="Practitioner:",
)

body(
    doc,
    "The paper also addresses the Saudi-Iran diplomatic normalisation scenario "
    "as a structural uncertainty that could reprice resilience investments made "
    "during the 2026 crisis — extending the practitioner message beyond "
    "'prepare for the last crisis' to 'prepare for a range of plausible futures.'",
)


# ══════════════════════════════════════════════════════════════════════════════
# Apply Times New Roman 11pt to all body text runs
# ══════════════════════════════════════════════════════════════════════════════
for p in doc.paragraphs:
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = (
    "/mnt/c/Users/User/Documents/GitHub/claude-squad/FRAGILE_FLOWS_Proposal_Group4.docx"
)
doc.save(out_path)
print(f"Saved: {out_path}")
