"""
Surgical fix to FlexCap_Case_Analysis_v11.docx — compliance push to 9-10/10.

Changes:
1. Section 2.5 heading: label "VC/PE Governance Principles" explicitly
2. Section 2.5 intro paragraph: reframe as VC/PE Governance concept application
3. Section 3.5 "Why This Addresses Root Cause": name the VC/PE Governance principle
   behind each mechanism (info rights cap, lock-up, ROFR)
4. Fix duplicate ROFR paragraph in 3.5 bullet
5. Fix missing bold on "Root Cause:" labels in 2.1, 2.2, 2.3
6. Fix "Root Cause:" label on 2.5 (same pattern as others)
7. Fix "Solution:" label bold on 3.1 (matching 3.2-3.5)
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import copy

SRC = "/mnt/c/Users/User/Desktop/FlexCap_Case_Analysis_v11.docx"
DST = "/mnt/c/Users/User/Desktop/FlexCap_Case_Analysis_v12.docx"

doc = Document(SRC)

# ─── FIX 1 & 2: Section 2.5 — relabel heading, reframe intro ────────────────

# Find the 2.5 heading paragraph
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text.startswith("2.5") and "SinoLife" in text:
        # Fix heading to explicitly include VC/PE Governance Principles
        for run in para.runs:
            if "2.5" in run.text:
                run.text = run.text.replace(
                    "2.5  SinoLife Strategic Distortion — Strategic vs. Financial Investor Risk",
                    "2.5  SinoLife Strategic Distortion — VC/PE Governance Principles Failure",
                )
                # Also ensure "Strategic vs. Financial Investor Risk" is removed from heading
                if "Strategic vs. Financial Investor Risk" in run.text:
                    run.text = run.text.replace(
                        " — Strategic vs. Financial Investor Risk", ""
                    )
        # If run.text didn't contain full string, rebuild
        if (
            "2.5" in para.text
            and "VC/PE Governance Principles Failure" not in para.text
        ):
            full = para.text
            para.clear()
            run = para.add_run(
                "2.5  SinoLife Strategic Distortion — VC/PE Governance Principles Failure"
            )
            run.bold = True
            run.font.size = Pt(11)
        print(f"Fix 1: 2.5 heading updated at para {i}")
        break

# ─── FIX 2: Section 2.5 intro — replace with VC/PE Governance framing ─────────

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text.startswith(
        "Strategic investor misalignment and revenue concentration risk distorting valuationRoot Cause:"
    ):
        # This is the 2.5 root cause paragraph — replace it entirely
        para.clear()
        # Rebuild with proper bold separators
        run_label = para.add_run(
            "Strategic investor misalignment and revenue concentration risk distorting valuation"
        )
        run_label.bold = False

        run_rc = para.add_run("Root Cause: ")
        run_rc.bold = True

        run_body = para.add_run(
            "SinoLife invested at $200M pre-money (vs. $500M implied post-money). "
            "VC/PE Governance Principles (Session 6 / Investment Governance) teach that strategic "
            "investors who simultaneously hold board seats, information rights, and commercial "
            "partnerships create compounded governance conflicts — the strategic agenda can override "
            "financial investor exit objectives, and information rights can be weaponized as competitive "
            "intelligence. Three structural tensions follow: (1) SinoLife's commercial agenda as a "
            "hospital network operator may diverge from Company X's independent growth path; "
            "(2) SinoLife's $200M → $500M step-up creates an asymmetric incentive to exit before "
            "FlexCap's governance rights materialize; and (3) SinoLife's information rights, if "
            "uncapped, allow a strategic competitor to monitor Company X's hospital partnership "
            "pipeline in real time — a direct conflict of interest that standard VC/PE governance "
            "contracts address through information rights separation, share lock-ups, and ROFR "
            "obligations on secondary transfers."
        )
        run_body.bold = False
        print(f"Fix 2: 2.5 Root Cause paragraph replaced at para {i}")
        break

# ─── FIX 3: Section 3.5 "Why This Addresses Root Cause" — name VC/PE principles ─

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text == "Why This Addresses Root Cause:" or text.startswith(
        "Why This Addresses Root Cause:"
    ):
        # Find the following paragraph with the actual body text
        next_i = i + 1
        if next_i < len(doc.paragraphs):
            next_para = doc.paragraphs[next_i]
            # Replace with VC/PE Governance-framed version
            next_para.clear()
            run = next_para.add_run(
                "VC/PE Governance Principles teach that strategic investors with board or observer "
                "seats, uncapped information rights, and no transfer restrictions create three "
                "distinct agency risks that standard financial covenants cannot resolve. First, "
                "information rights separation prevents SinoLife from using investor disclosure "
                "channels to gather competitive intelligence on Company X's hospital partnerships — "
                "a principle directly applicable when the strategic investor operates in the same "
                "market. Second, a 24-month share lock-up prevents SinoLife from flipping its equity "
                "at the $200M → $500M step-up before FlexCap's value-creation milestones are met — "
                "aligning holding period with investment horizon. Third, the ROFR obligation ensures "
                "that any SinoLife secondary transfer triggers a matching right, preventing a "
                "far-term sale to a strategic acquirer who would convert Company X into a "
                "captured-channel asset. These three instruments — information rights cap, lock-up, "
                "and ROFR — are the standard VC/PE governance toolkit for precisely this conflict "
                "type. SinoLife's misaligned incentives cannot be resolved by earn-outs or "
                "milestone pricing alone; only a renegotiated governance structure addresses the "
                "root conflict."
            )
            run.bold = False
            print(f"Fix 3: 3.5 'Why This Addresses' replaced at para {next_i}")
        break

# ─── FIX 4: Fix duplicate ROFR paragraph in 3.5 bullet ──────────────────────

# The ROFR bullet at line 90 is duplicated — two identical sentences concatenated.
# Find it and fix.
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if "SinoLife ROFR obligation" in text and "Prev" in text:
        # This is the duplicate — fix it to show only one clean version
        para.clear()
        run = para.add_run(
            "SinoLife ROFR obligation: SinoLife's shares are subject to a "
            "Right-of-First-Refusal obligation — before any transfer, sale, or assignment "
            "of SinoLife's shares to a third party, Company X (or its designee) has the "
            "right to match any proposed transaction. This prevents SinoLife from flipping "
            "equity shortly after Series B at a profit — realistic given the $200M → $500M "
            "step-up — and aligns SinoLife's incentives with Company X's long-term value "
            "creation."
        )
        run.bold = False
        print(f"Fix 4: Duplicate ROFR paragraph cleaned at para {i}")
        break

# ─── FIX 5 & 6: Fix Root Cause: label bold consistency in 2.1, 2.2, 2.3, 2.5 ─

# Find all "Root Cause:" patterns and ensure label is bold + rest is normal
for para in doc.paragraphs:
    text = para.text
    if "Root Cause:" in text:
        # Reconstruct: split on "Root Cause:"
        parts = text.split("Root Cause:")
        if len(parts) == 2:
            before = parts[0].rstrip()
            after = parts[1].lstrip() if len(parts) > 1 else ""
            # Clear and rebuild
            para.clear()
            r1 = para.add_run(before)
            r1.bold = False
            r2 = para.add_run("Root Cause: ")
            r2.bold = True
            r3 = para.add_run(after)
            r3.bold = False

# ─── FIX 7: Fix "Solution:" label bold consistency in 3.1 ────────────────────

for para in doc.paragraphs:
    text = para.text
    if "Solution:" in text and text.startswith("Renegotiate valuation"):
        parts = text.split("Solution:")
        if len(parts) == 2:
            before = parts[0].rstrip()
            after = parts[1].lstrip()
            para.clear()
            r1 = para.add_run(before)
            r1.bold = False
            r2 = para.add_run("Solution: ")
            r2.bold = True
            r3 = para.add_run(after)
            r3.bold = False
            print(f"Fix 7: 3.1 Solution label fixed")
            break

# Save as v12
doc.save(DST)
print(f"\nSaved → {DST}")
print("All 7 fixes applied.")
