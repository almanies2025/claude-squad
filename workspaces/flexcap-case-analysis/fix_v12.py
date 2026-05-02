"""
Surgical targeted fix for FlexCap_Case_Analysis_v12.docx
Fixes remaining issues from fix_v11.py:
- Para [49]: 2.5 Root Cause label merged with previous sentence
- Para [50]: orphaned duplicate first-sentence of new 2.5 text
- Para [99]: "Why This Addresses" merged with "Root Cause:" (no space)
- Para [100]: old "Why This Addresses" text not replaced
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

SRC = "/mnt/c/Users/User/Desktop/FlexCap_Case_Analysis_v12.docx"
DST = "/mnt/c/Users/User/Desktop/FlexCap_Case_Analysis_v13.docx"

doc = Document(SRC)


# ── Helper: clear all runs from a paragraph and add one plain text run ─────────
def replace_para_text(para, new_text):
    for run in para.runs:
        run.text = ""
    # Clear any existing XML children that aren't runs (e.g., bold marks)
    for child in list(para._p):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag not in ("run", "tbl", "tc"):
            para._p.remove(child)
    # Add single clean run
    run = para.add_run(new_text)
    run.bold = False
    run.font.size = Pt(11)


# ── Helper: clear runs and set label+body with bold label ────────────────────
def rebuild_para(para, label, body):
    """label is bold, body is normal"""
    for run in para.runs:
        run.text = ""
    r_label = para.add_run(label)
    r_label.bold = True
    r_label.font.size = Pt(11)
    r_body = para.add_run(body)
    r_body.bold = False
    r_body.font.size = Pt(11)


# ── Issue 1: Para [49] — "Root Cause:" merged with previous sentence ─────────
# Fix: clear and rebuild with proper "Root Cause: " label + body
for para in doc.paragraphs:
    if (
        "Strategic investor misalignment and revenue concentration risk distorting valuationRoot Cause:"
        in para.text
    ):
        rebuild_para(
            para,
            "Root Cause: ",
            "SinoLife invested at $200M pre-money (vs. $500M implied post-money). "
            "VC/PE Governance Principles (Session 6 / Investment Governance) teach that strategic "
            "investors who simultaneously hold board seats, information rights, and commercial "
            "partnerships create compounded governance conflicts — the strategic agenda can override "
            "financial investor exit objectives, and information rights can be weaponized as competitive "
            "intelligence. Three structural tensions follow: (1) SinoLife's commercial agenda as a "
            "hospital network operator may diverge from Company X's independent growth path; "
            "(2) SinoLife's $200M → $500M step-up creates an asymmetric incentive to exit before "
            "FlexCap's governance rights materialize; and (3) SinoLife's information rights, if "
            "uncapped, allow a strategic competitor to monitor Company X's hospital partnership "
            "pipeline in real time — a direct conflict of interest that standard VC/PE governance "
            "contracts address through information rights separation, share lock-ups, and ROFR "
            "obligations on secondary transfers.",
        )
        print("Fix 1: 2.5 Root Cause paragraph rebuilt")
        break

# ── Issue 2: Para [50] — orphaned duplicate sentence (new text starting mid-sentence) ─
# This paragraph is the first sentence of the new text — now redundant with para [49]
for para in doc.paragraphs:
    if (
        para.text.strip()
        == "SinoLife invested at $200M pre-money (vs. $500M implied post-money), creating three structural tensions:"
    ):
        replace_para_text(para, "")
        print("Fix 2: Orphaned para [50] cleared")
        break

# ── Issue 3: Para [99] — "Why This Addresses" merged with "Root Cause:" ────────
for para in doc.paragraphs:
    if "Why This AddressesRoot Cause:" in para.text:
        rebuild_para(
            para,
            "Why This Addresses Root Cause: ",
            "VC/PE Governance Principles teach that strategic investors with board or observer "
            "seats, uncapped information rights, and no transfer restrictions create three "
            "distinct agency risks that standard financial covenants cannot resolve. First, "
            "information rights separation prevents SinoLife from using investor disclosure "
            "channels to gather competitive intelligence on Company X's hospital partnerships — "
            "a principle directly applicable when the strategic investor operates in the same "
            "market. Second, a 24-month share lock-up prevents SinoLife from flipping its equity "
            "at the $200M → $500M step-up before FlexCap's value-creation milestones are met — "
            "aligning holding period with investment horizon. Third, the ROFR obligation ensures "
            "that any SinoLife secondary transfer triggers a matching right, preventing a "
            "far-term sale to a strategic acquirer who would convert Company X into a "
            "captured-channel asset. These three instruments — information rights cap, lock-up, "
            "and ROFR — are the standard VC/PE governance toolkit for precisely this conflict "
            "type. SinoLife's misaligned incentives cannot be resolved by earn-outs or "
            "milestone pricing alone; only a renegotiated governance structure addresses the "
            "root conflict.",
        )
        print("Fix 3: 3.5 'Why This Addresses Root Cause' paragraph rebuilt")
        break

# ── Issue 4: Para [100] — old text still there after Fix 3 ────────────────────
# (Fix 3 replaced para [99], so para [100] with old text is now orphaned)
for para in doc.paragraphs:
    if (
        "SinoLife's misaligned incentives are a structural problem that earn-outs"
        in para.text
    ):
        replace_para_text(para, "")
        print("Fix 4: Old 3.5 'Why This Addresses' body paragraph cleared")
        break

# ── Issue 5: Root Cause label bold fix for 2.1, 2.2, 2.3, 2.4 ─────────────
# Pattern: "sentence fragmentRoot Cause: body text" — need to split and bold properly
for para in doc.paragraphs:
    t = para.text
    marker = "Root Cause:"
    if marker in t:
        # Skip if already processed (already has "Root Cause: " as its own label)
        # Pattern to detect: looks like "Some sentence textRoot Cause: rest"
        # vs properly formatted: has "Root Cause: " preceded by newline or space+bold
        idx = t.index(marker)
        before = t[:idx]
        after = t[idx + len(marker) :].lstrip()
        # If 'before' is a sentence fragment (doesn't end with a period or is very long),
        # it's the merged pattern
        if before and not before.endswith(".") and len(before) > 10:
            rebuild_para(para, "Root Cause: ", after)
            print(f"Fix 5: Rebuilt merged Root Cause label: {before[:40]}...")
        # If it starts with "Root Cause:" already as label, skip

# ── Issue 6: "Why This AddressesRoot Cause:" in any other paragraph ────────────
for para in doc.paragraphs:
    if "Why This AddressesRoot Cause:" in para.text:
        rebuild_para(
            para,
            "Why This Addresses Root Cause: ",
            "VC/PE Governance Principles teach that strategic investors with board or observer "
            "seats, uncapped information rights, and no transfer restrictions create three "
            "distinct agency risks that standard financial covenants cannot resolve. First, "
            "information rights separation prevents SinoLife from using investor disclosure "
            "channels to gather competitive intelligence on Company X's hospital partnerships — "
            "a principle directly applicable when the strategic investor operates in the same "
            "market. Second, a 24-month share lock-up prevents SinoLife from flipping its equity "
            "at the $200M → $500M step-up before FlexCap's value-creation milestones are met — "
            "aligning holding period with investment horizon. Third, the ROFR obligation ensures "
            "that any SinoLife secondary transfer triggers a matching right, preventing a "
            "far-term sale to a strategic acquirer who would convert Company X into a "
            "captured-channel asset. These three instruments — information rights cap, lock-up, "
            "and ROFR — are the standard VC/PE governance toolkit for precisely this conflict "
            "type. SinoLife's misaligned incentives cannot be resolved by earn-outs or "
            "milestone pricing alone; only a renegotiated governance structure addresses the "
            "root conflict.",
        )
        print("Fix 6: Additional 'Why This AddressesRoot Cause' cleaned")
        break

doc.save(DST)
print(f"\nSaved → {DST}")
