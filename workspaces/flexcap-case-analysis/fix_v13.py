"""
FlexCap_Case_Analysis_v13.docx — targeted fixes only.

Changes from v12:
1. Fix 4 regressions: "Why This AddressesRoot Cause:" → "Why This Addresses Root Cause:"
   (paras [74], [83], [91], [99]) — Fix 5 accidentally removed the space.
2. Reframe 2.5 Root Cause body (para [49]) to explicitly cite VC/PE Governance Principles
   as the course concept, showing analytical application not just description.
3. Reframe 3.5 "Why This Addresses" body (para [100]) with explicit VC/PE Governance
   principle names behind each mechanism — this is the main compliance gap.

Everything else in v12 stays as-is.
"""

from docx import Document
from docx.shared import Pt

SRC = "/mnt/c/Users/User/Desktop/FlexCap_Case_Analysis_v12.docx"
DST = "/mnt/c/Users/User/Desktop/FlexCap_Case_Analysis_v13.docx"

doc = Document(SRC)

changed = []

# ── FIX 1: Restore space in "Why This AddressesRoot Cause:" labels ─────────────
# Fix 5 removed the space — restore it. These labels appear in 3.2, 3.3, 3.4, 3.5.
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == "Why This AddressesRoot Cause:":
        for run in para.runs:
            if run.text == "Why This AddressesRoot Cause:":
                run.text = "Why This Addresses Root Cause: "
                changed.append(f"Fix 1: Restored space in para [{i}]")
                break

# ── FIX 2: Reframe 2.5 Root Cause body with VC/PE Governance concept ───────────
# Para [49] — the Root Cause body for SinoLife.
# Replace the current text (v11 content) with an explicit VC/PE Governance framing.
for i, para in enumerate(doc.paragraphs):
    if i == 49 and "Strategic investor misalignment" in para.text:
        # Replace the run(s) with properly framed content
        # Clear all runs
        for run in para.runs:
            run.text = ""
        # Add label run (bold)
        r_label = para.add_run("Root Cause: ")
        r_label.bold = True
        r_label.font.size = Pt(11)
        # Add body run
        r_body = para.add_run(
            "SinoLife invested at $200M pre-money (vs. $500M implied post-money). "
            "VC/PE Governance Principles (Session 6) teach that strategic investors who "
            "simultaneously hold equity stakes, board/observer seats, and commercial partnerships "
            "create compounded agency conflicts — the strategic agenda can override financial "
            "investor exit objectives, and information rights can be weaponized as competitive "
            "intelligence rather than governance oversight. Three structural tensions follow: "
            "(1) SinoLife's commercial agenda as a hospital network operator may diverge from "
            "Company X's independent growth path; (2) SinoLife's $200M → $500M step-up creates "
            "an asymmetric incentive to exit before FlexCap's governance rights are operative; "
            "and (3) SinoLife's information rights, if uncapped, allow a strategic competitor "
            "to monitor Company X's hospital partnership pipeline in real time — a conflict "
            "that standard VC/PE governance contracts address through information rights "
            "separation, share lock-ups, and ROFR obligations on secondary transfers. "
            "The SinoLife strategic investment is therefore not merely a valuation tension "
            "— it is a VC/PE governance failure requiring a governance-level solution."
        )
        r_body.bold = False
        r_body.font.size = Pt(11)
        changed.append(f"Fix 2: 2.5 Root Cause reframe at para [{i}]")
        break

# ── FIX 3: Reframe 3.5 "Why This Addresses Root Cause" body ──────────────────
# Para [100] — replace the v11/v12 generic text with VC/PE Governance-framed version.
for i, para in enumerate(doc.paragraphs):
    if (
        i == 100
        and "SinoLife's misaligned incentives are a structural problem" in para.text
    ):
        for run in para.runs:
            run.text = ""
        r = para.add_run(
            "VC/PE Governance Principles identify three distinct agency risks in deals with "
            "strategic investors holding simultaneous equity stakes, board access, and commercial "
            "relationships — each requiring a specific governance instrument. First, information "
            "rights separation (capping SinoLife to quarterly financial updates, not operational "
            "data) is the standard VC/PE remedy when a strategic investor operates in the same "
            "market — it prevents investor disclosure channels from becoming a competitive "
            "intelligence pipeline. Second, a 24-month share lock-up is the VC/PE mechanism "
            "for aligning a strategic investor's holding period with the financial investor's "
            "value-creation horizon — preventing a short-term flip at the $200M → $500M step-up "
            "before FlexCap's milestones are tested. Third, a ROFR obligation on any SinoLife "
            "secondary transfer is the VC/PE instrument that prevents sale to a strategic acquirer "
            "who would convert Company X into a captured-channel asset, preserving Company X's "
            "independence for FlexCap's exit optionality. These three instruments — information "
            "rights cap, lock-up, and ROFR — are the standard VC/PE governance toolkit for this "
            "conflict type. SinoLife's misaligned incentives cannot be resolved by earn-outs "
            "or milestone pricing alone; only a renegotiated governance structure addresses "
            "the root conflict, as VC/PE Governance Principles prescribe."
        )
        r.bold = False
        r.font.size = Pt(11)
        changed.append(
            f"Fix 3: 3.5 'Why This Addresses Root Cause' body replaced at para [{i}]"
        )
        break

doc.save(DST)

print(f"Saved → {DST}")
print(f"\n{len(changed)} targeted fix(es) applied:")
for c in changed:
    print(f"  {c}")
